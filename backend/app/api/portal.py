import html
import json
import os
import re
import shutil
import uuid
import zipfile
from datetime import datetime
from pathlib import Path
from typing import Any
from urllib.parse import quote

from fastapi import APIRouter, File, HTTPException, Request, UploadFile
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session, selectinload

from app import models
from app.api.deps import CurrentUser, DbSession
from app.api.router_utils import apply_filters, paginate
from app.core.config import get_settings
from app.serializers import serialize_file_review_checklist, serialize_review_rule
from app.services import AIService

try:
    import docx
except ImportError:  # pragma: no cover - dependency is declared in requirements
    docx = None

try:
    import fitz
except ImportError:  # pragma: no cover - dependency is declared in requirements
    fitz = None


router = APIRouter(prefix="/portal", tags=["portal"])
settings = get_settings()

PORTAL_DOCUMENT_ROOT = settings.media_root / "portal" / "documents"
PORTAL_ALLOWED_EXTENSIONS = {".docx", ".pdf"}
PORTAL_MAX_UPLOAD_SIZE = int(os.getenv("PORTAL_MAX_UPLOAD_SIZE", str(20 * 1024 * 1024)))
PORTAL_DOCX_CHARS_PER_PAGE = int(os.getenv("PORTAL_DOCX_CHARS_PER_PAGE", "1800"))
PORTAL_DOCUMENT_ID_RE = re.compile(r"^[a-f0-9]{32}$")

DOCUMENT_OVERVIEW_SYSTEM_PROMPT = (
    "你是文档总结概览助手。请根据用户提供的文档全文生成文档概览。"
    "要求：只输出中文概览正文，不要标题、编号、Markdown或解释说明；"
    "概览必须聚焦文档主题、核心内容、关键事项和结论，控制在100到400字。"
)

PORTAL_RULE_REVIEW_SYSTEM_PROMPT = (
    "你是企业文件审查专家。请严格依据用户提供的审查规则，对文档页分片进行逐页审查。"
    "只输出JSON，不要输出Markdown代码块或额外解释。"
)
PORTAL_REVIEW_CHUNK_SIZE = int(os.getenv("PORTAL_REVIEW_CHUNK_SIZE", "1200"))


def sanitize_filename(filename: str) -> str:
    name = Path(filename or "").name.strip()
    if not name:
        return "未命名文档"
    return re.sub(r"[\\/:*?\"<>|]+", "_", name)[:180]


def count_preview_words(text: str) -> int:
    return len(re.findall(r"[\u4e00-\u9fff]|[A-Za-z0-9]+(?:[-_][A-Za-z0-9]+)*", text or ""))


def validate_portal_document_id(document_id: str) -> None:
    if not PORTAL_DOCUMENT_ID_RE.match(document_id):
        raise HTTPException(status_code=404, detail="文档不存在")


def portal_document_dir(document_id: str) -> Path:
    validate_portal_document_id(document_id)
    return PORTAL_DOCUMENT_ROOT / document_id


def portal_metadata_path(document_id: str) -> Path:
    return portal_document_dir(document_id) / "metadata.json"


def load_portal_metadata(document_id: str) -> dict[str, Any]:
    metadata_path = portal_metadata_path(document_id)
    if not metadata_path.exists():
        raise HTTPException(status_code=404, detail="文档不存在")
    try:
        return json.loads(metadata_path.read_text(encoding="utf-8"))
    except json.JSONDecodeError as exc:
        raise HTTPException(status_code=500, detail="文档元数据损坏") from exc


def save_portal_metadata(metadata: dict[str, Any]) -> None:
    PORTAL_DOCUMENT_ROOT.mkdir(parents=True, exist_ok=True)
    metadata_path = portal_metadata_path(metadata["id"])
    metadata_path.write_text(json.dumps(metadata, ensure_ascii=False, indent=2), encoding="utf-8")


def validate_uploaded_file_signature(file_path: Path, extension: str) -> None:
    if extension == ".pdf":
        with file_path.open("rb") as source:
            if source.read(4) != b"%PDF":
                raise HTTPException(status_code=400, detail="文件内容不是有效的 PDF")
        return
    if extension == ".docx":
        if not zipfile.is_zipfile(file_path):
            raise HTTPException(status_code=400, detail="文件内容不是有效的 DOCX")
        with zipfile.ZipFile(file_path) as archive:
            names = set(archive.namelist())
            if "[Content_Types].xml" not in names or "word/document.xml" not in names:
                raise HTTPException(status_code=400, detail="文件内容不是有效的 DOCX")


def docx_table_to_block(table: Any) -> dict[str, str]:
    rows: list[str] = []
    html_rows: list[str] = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        if any(cells):
            rows.append(" | ".join(cells))
        html_cells = "".join(f"<td>{html.escape(cell)}</td>" for cell in cells)
        html_rows.append(f"<tr>{html_cells}</tr>")
    text = "\n".join(rows)
    return {
        "text": text,
        "html": f'<table class="docx-preview-table"><tbody>{"".join(html_rows)}</tbody></table>',
    }


def extract_docx_preview(file_path: Path) -> dict[str, Any]:
    if docx is None:
        raise HTTPException(status_code=500, detail="服务端缺少 python-docx 依赖")

    document = docx.Document(file_path)
    blocks: list[dict[str, str]] = []
    text_parts: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        paragraph_xml = paragraph._element.xml
        has_page_break = 'w:type="page"' in paragraph_xml or "lastRenderedPageBreak" in paragraph_xml
        if not text:
            if has_page_break:
                blocks.append({"text": "", "html": "", "page_break": "1"})
            continue
        text_parts.append(text)
        style_name = (getattr(paragraph.style, "name", "") or "").lower()
        tag = "h2" if "heading" in style_name or "标题" in style_name else "p"
        blocks.append({"text": text, "html": f"<{tag}>{html.escape(text)}</{tag}>"})
        if has_page_break:
            blocks.append({"text": "", "html": "", "page_break": "1"})

    for table in document.tables:
        block = docx_table_to_block(table)
        if block["text"]:
            text_parts.append(block["text"])
            blocks.append(block)

    pages: list[dict[str, Any]] = []
    current_html: list[str] = []
    current_text: list[str] = []
    current_chars = 0
    for block in blocks:
        if block.get("page_break"):
            if current_html:
                pages.append(
                    {
                        "page_number": len(pages) + 1,
                        "type": "docx_html",
                        "html": "".join(current_html),
                        "text": "\n".join(current_text),
                    }
                )
                current_html = []
                current_text = []
                current_chars = 0
            continue

        block_len = len(block["text"])
        if current_html and current_chars + block_len > PORTAL_DOCX_CHARS_PER_PAGE:
            pages.append(
                {
                    "page_number": len(pages) + 1,
                    "type": "docx_html",
                    "html": "".join(current_html),
                    "text": "\n".join(current_text),
                }
            )
            current_html = []
            current_text = []
            current_chars = 0

        current_html.append(block["html"])
        current_text.append(block["text"])
        current_chars += block_len

    if current_html:
        pages.append(
            {
                "page_number": len(pages) + 1,
                "type": "docx_html",
                "html": "".join(current_html),
                "text": "\n".join(current_text),
            }
        )
    if not pages:
        pages.append({"page_number": 1, "type": "docx_html", "html": "<p>该文档没有可预览文本。</p>", "text": ""})

    full_text = "\n".join(text_parts)
    try:
        title = (document.core_properties.title or "").strip()
    except Exception:
        title = ""
    if not title:
        title = next((part for part in text_parts if part), "")

    return {
        "title": title[:120],
        "page_count": len(pages),
        "word_count": count_preview_words(full_text),
        "pages": pages,
    }


def extract_pdf_preview(document_id: str, file_path: Path) -> dict[str, Any]:
    if fitz is None:
        raise HTTPException(status_code=500, detail="服务端缺少 PyMuPDF 依赖")

    text_parts: list[str] = []
    pages: list[dict[str, Any]] = []
    with fitz.open(file_path) as pdf:
        metadata = pdf.metadata or {}
        for page_index in range(len(pdf)):
            page = pdf[page_index]
            text = page.get_text("text").strip()
            text_parts.append(text)
            pages.append(
                {
                    "page_number": page_index + 1,
                    "type": "pdf_image",
                    "image_url": f"/api/portal/documents/{document_id}/pages/{page_index + 1}.png",
                    "text": text,
                }
            )
        title = (metadata.get("title") or "").strip()
        if not title:
            first_page_text = text_parts[0] if text_parts else ""
            first_lines = [line.strip() for line in first_page_text.splitlines() if line.strip()]
            title = first_lines[0] if first_lines else ""

    if not pages:
        pages.append(
            {
                "page_number": 1,
                "type": "pdf_image",
                "image_url": f"/api/portal/documents/{document_id}/pages/1.png",
                "text": "",
            }
        )

    full_text = "\n".join(text_parts)
    return {
        "title": title[:120],
        "page_count": len(pages),
        "word_count": count_preview_words(full_text),
        "pages": pages,
    }


def build_portal_preview(document_id: str, file_path: Path, extension: str) -> dict[str, Any]:
    if extension == ".pdf":
        return extract_pdf_preview(document_id, file_path)
    if extension == ".docx":
        return extract_docx_preview(file_path)
    raise HTTPException(status_code=400, detail="仅支持 DOCX、PDF 文件")


def portal_url(request: Request, path: str) -> str:
    base_url = str(request.base_url).rstrip("/")
    return f"{base_url}{path}"


def serialize_portal_document(request: Request, metadata: dict[str, Any]) -> dict[str, Any]:
    document_id = metadata["id"]
    file_url = f"/api/portal/documents/{document_id}/file/"
    preview_url = f"/api/portal/documents/{document_id}/preview/"
    pages: list[dict[str, Any]] = []
    for page in metadata.get("pages", []):
        page_data = page.copy()
        if page_data.get("image_url"):
            page_data["absolute_image_url"] = portal_url(request, page_data["image_url"])
        pages.append(page_data)
    return {
        "id": document_id,
        "name": metadata["original_name"],
        "title": metadata.get("title") or metadata["original_name"],
        "extension": metadata["extension"],
        "content_type": metadata["content_type"],
        "size": metadata["size"],
        "page_count": metadata["page_count"],
        "word_count": metadata["word_count"],
        "created_at": metadata["created_at"],
        "file_url": file_url,
        "preview_url": preview_url,
        "absolute_file_url": portal_url(request, file_url),
        "absolute_preview_url": portal_url(request, preview_url),
        "pages": pages,
    }


def normalize_document_overview(text: str) -> str:
    overview = re.sub(r"\s+", " ", text or "").strip()
    if len(overview) <= 400:
        return overview
    clipped = overview[:400]
    for separator in ("。", "！", "？", "；"):
        index = clipped.rfind(separator)
        if index >= 100:
            return clipped[: index + 1].strip()
    return clipped.strip()


def fallback_document_overview(content: str, document_name: str = "") -> str:
    text = re.sub(r"\s+", " ", content or "").strip()
    if not text:
        return "文档暂无可提取文本，无法生成详细概览。"

    sentences = [part.strip() for part in re.split(r"(?<=[。！？；])", text) if part.strip()]
    if not sentences:
        sentences = [text]

    selected: list[str] = []
    total_len = 0
    for sentence in sentences:
        if total_len >= 280:
            break
        selected.append(sentence)
        total_len += len(sentence)

    summary = "".join(selected).strip()
    if document_name:
        summary = f"{document_name}主要内容包括：{summary}"
    return normalize_document_overview(summary)


def split_review_chunks(text: str, chunk_size: int = PORTAL_REVIEW_CHUNK_SIZE) -> list[dict[str, Any]]:
    normalized = re.sub(r"\n{3,}", "\n\n", text or "").strip()
    if not normalized:
        return []

    chunks: list[dict[str, Any]] = []
    cursor = 0
    while cursor < len(normalized):
        chunk = normalized[cursor : cursor + chunk_size].strip()
        if chunk:
            chunks.append({"index": len(chunks) + 1, "content": chunk})
        cursor += chunk_size
    return chunks


def compact_rule_for_prompt(rule: dict[str, Any]) -> dict[str, Any]:
    return {
        "rule_code": rule.get("rule_code") or "",
        "rule_name": rule.get("rule_name") or "",
        "rule_type": rule.get("rule_type") or "",
        "category": rule.get("category") or "",
        "risk_level": rule.get("risk_level") or "",
        "description": rule.get("description") or "",
        "legal_basis": rule.get("legal_basis") or "",
        "rule_content": rule.get("rule_content") or {},
    }


def extract_json_object(text: str) -> dict[str, Any]:
    cleaned = (text or "").strip()
    cleaned = re.sub(r"^```(?:json)?\s*", "", cleaned)
    cleaned = re.sub(r"\s*```$", "", cleaned)
    try:
        payload = json.loads(cleaned)
        return payload if isinstance(payload, dict) else {"raw": payload}
    except json.JSONDecodeError:
        start = cleaned.find("{")
        end = cleaned.rfind("}")
        if start >= 0 and end > start:
            try:
                payload = json.loads(cleaned[start : end + 1])
                return payload if isinstance(payload, dict) else {"raw": payload}
            except json.JSONDecodeError:
                pass
    return {"summary": cleaned[:1000], "issues": [], "raw_response": cleaned}


def normalize_review_page_result(page_number: int, chunks: list[dict[str, Any]], response_text: str) -> dict[str, Any]:
    payload = extract_json_object(response_text)
    issues = payload.get("issues") if isinstance(payload.get("issues"), list) else []
    normalized_issues: list[dict[str, Any]] = []
    for issue in issues:
        if not isinstance(issue, dict):
            continue
        normalized_issues.append(
            {
                "rule_code": issue.get("rule_code") or "",
                "rule_name": issue.get("rule_name") or "",
                "risk_level": issue.get("risk_level") or "",
                "clause": issue.get("clause") or issue.get("evidence") or "",
                "issue": issue.get("issue") or issue.get("description") or "",
                "suggestion": issue.get("suggestion") or "",
                "legal_basis": issue.get("legal_basis") or "",
                "chunk_index": issue.get("chunk_index"),
            }
        )

    return {
        "page_number": page_number,
        "chunk_count": len(chunks),
        "summary": payload.get("summary") or "",
        "risk_level": payload.get("risk_level") or ("high" if any(item.get("risk_level") == "high" for item in normalized_issues) else "medium" if normalized_issues else "low"),
        "issues": normalized_issues,
        "raw_response": payload.get("raw_response") or response_text,
    }


def build_rule_review_messages(document_name: str, page: dict[str, Any], chunks: list[dict[str, Any]], rules: list[dict[str, Any]]) -> list[dict[str, str]]:
    page_number = page.get("page_number")
    payload = {
        "document_name": document_name,
        "page_number": page_number,
        "review_rules": rules,
        "page_chunks": chunks,
    }
    output_schema = {
        "page_number": page_number,
        "summary": "本页审查结论摘要",
        "risk_level": "low/medium/high",
        "issues": [
            {
                "rule_code": "命中的规则编码",
                "rule_name": "命中的规则名称",
                "risk_level": "low/medium/high",
                "chunk_index": 1,
                "clause": "命中的原文证据或摘要",
                "issue": "问题描述",
                "legal_basis": "法律依据或规则依据",
                "suggestion": "修改建议",
            }
        ],
    }
    user_content = (
        "请依据审查规则审查以下文档页分片。"
        "仅识别与规则直接相关的问题；没有问题时 issues 返回空数组，risk_level 返回 low。"
        "返回JSON格式必须符合 output_schema。\n\n"
        f"input:\n{json.dumps(payload, ensure_ascii=False)}\n\n"
        f"output_schema:\n{json.dumps(output_schema, ensure_ascii=False)}"
    )
    return [
        {"role": "system", "content": PORTAL_RULE_REVIEW_SYSTEM_PROMPT},
        {"role": "user", "content": user_content},
    ]


def get_default_ai_config(db: Session) -> models.AIModelConfig | None:
    return (
        db.query(models.AIModelConfig)
        .filter(models.AIModelConfig.is_default.is_(True), models.AIModelConfig.is_active.is_(True))
        .order_by(models.AIModelConfig.id.desc())
        .first()
    ) or (
        db.query(models.AIModelConfig)
        .filter(models.AIModelConfig.is_active.is_(True))
        .order_by(models.AIModelConfig.is_default.desc(), models.AIModelConfig.id.desc())
        .first()
    )


def portal_checklist_query(db: Session):
    return db.query(models.FileReviewChecklist).options(
        selectinload(models.FileReviewChecklist.created_by),
        selectinload(models.FileReviewChecklist.updated_by),
        selectinload(models.FileReviewChecklist.rule_links)
        .selectinload(models.FileReviewChecklistRule.rule)
        .selectinload(models.ReviewRule.created_by),
    )


def serialize_portal_review_record(record: models.PortalFileReviewRecord, include_result: bool = False) -> dict[str, Any]:
    summary = record.summary or {}
    payload = {
        "id": record.id,
        "document_id": record.document_id,
        "document_name": record.document_name,
        "document_extension": record.document_extension,
        "document_size": record.document_size,
        "page_count": record.page_count,
        "word_count": record.word_count,
        "checklist_id": record.checklist_id,
        "checklist_name": record.checklist_name,
        "rule_count": record.rule_count,
        "status": record.status,
        "model": record.model,
        "summary": summary,
        "risk_level": summary.get("risk_level") or "",
        "issue_count": summary.get("issue_count") or 0,
        "error_message": record.error_message,
        "started_at": record.started_at.isoformat() if record.started_at else None,
        "completed_at": record.completed_at.isoformat() if record.completed_at else None,
        "created_at": record.created_at.isoformat() if record.created_at else None,
        "updated_at": record.updated_at.isoformat() if record.updated_at else None,
        "created_by": record.created_by.username if record.created_by else "",
    }
    if include_result:
        payload["pages"] = record.pages or []
        payload["request_payload"] = record.request_payload or {}
    return payload


def mark_portal_review_record_failed(
    db: Session,
    record: models.PortalFileReviewRecord | None,
    message: str,
) -> None:
    if not record:
        return
    record.status = "failed"
    record.error_message = str(message)[:2000]
    record.completed_at = datetime.now()
    db.add(record)
    db.commit()


@router.get("/knowledge/checklists/")
def list_portal_checklists(request: Request, db: DbSession):
    query = portal_checklist_query(db)
    query = apply_filters(
        query,
        models.FileReviewChecklist,
        request,
        filter_fields=["name"],
        search_fields=["name", "description"],
        ordering_fields=["updated_at", "created_at", "name"],
        default_ordering=["-updated_at"],
    )
    return paginate(query, request, serialize_file_review_checklist)


@router.get("/knowledge/checklists/{item_id}/")
def retrieve_portal_checklist(item_id: int, db: DbSession):
    item = portal_checklist_query(db).filter(
        models.FileReviewChecklist.id == item_id,
        models.FileReviewChecklist.is_deleted.is_(False),
    ).first()
    if not item:
        raise HTTPException(status_code=404, detail="资源不存在")
    return serialize_file_review_checklist(item)


@router.get("/knowledge/rules/")
def list_portal_rules(request: Request, db: DbSession):
    query = db.query(models.ReviewRule).options(selectinload(models.ReviewRule.created_by))
    query = apply_filters(
        query,
        models.ReviewRule,
        request,
        filter_fields=["rule_type", "industry", "category", "risk_level", "is_active", "rule_name", "rule_code"],
        search_fields=["rule_name", "rule_code", "description"],
        ordering_fields=["priority", "created_at"],
        default_ordering=["-priority", "-created_at"],
        aliases={"created_by": "created_by_id"},
    )
    query = query.filter(models.ReviewRule.is_active.is_(True))
    return paginate(query, request, serialize_review_rule)


@router.get("/knowledge/rules/{item_id}/")
def retrieve_portal_rule(item_id: int, db: DbSession):
    item = (
        db.query(models.ReviewRule)
        .options(selectinload(models.ReviewRule.created_by))
        .filter(
            models.ReviewRule.id == item_id,
            models.ReviewRule.is_deleted.is_(False),
            models.ReviewRule.is_active.is_(True),
        )
        .first()
    )
    if not item:
        raise HTTPException(status_code=404, detail="资源不存在")
    return serialize_review_rule(item)


@router.post("/ai/document-overview/")
@router.post("/ai/document-overview-chat/")
async def portal_document_overview(request: Request, db: DbSession):
    payload = await request.json()
    content = str(
        payload.get("content")
        or payload.get("document_text")
        or payload.get("message")
        or ""
    ).strip()
    if not content:
        raise HTTPException(status_code=400, detail="文档内容不能为空")

    document_name = str(payload.get("document_name") or payload.get("filename") or "").strip()
    user_content = "请为以下文档生成100到400字的概览。"
    if document_name:
        user_content += f"\n文档名称：{document_name}"
    user_content += f"\n\n文档全文：\n{content}"

    model = "fallback"
    try:
        config = get_default_ai_config(db)
        response_text = AIService(db, config=config).call(
            [
                {"role": "system", "content": DOCUMENT_OVERVIEW_SYSTEM_PROMPT},
                {"role": "user", "content": user_content},
            ],
            max_tokens=700,
            timeout=120,
        )
        model = config.default_model if config else model
    except Exception:
        response_text = fallback_document_overview(content, document_name)

    overview = normalize_document_overview(response_text)
    return {"success": True, "overview": overview, "response": overview, "model": model}


@router.get("/reviews/history/")
def list_portal_review_history(request: Request, db: DbSession, current_user: CurrentUser):
    query = db.query(models.PortalFileReviewRecord).options(selectinload(models.PortalFileReviewRecord.created_by))
    if not current_user.is_superuser:
        query = query.filter(models.PortalFileReviewRecord.created_by_id == current_user.id)
    query = apply_filters(
        query,
        models.PortalFileReviewRecord,
        request,
        filter_fields=["status", "checklist_id", "document_extension"],
        search_fields=["document_name", "checklist_name"],
        ordering_fields=["created_at", "completed_at", "document_name", "status"],
        default_ordering=["-created_at"],
    )
    return paginate(query, request, lambda record: serialize_portal_review_record(record))


@router.get("/reviews/history/{record_id}/")
def retrieve_portal_review_history(record_id: int, db: DbSession, current_user: CurrentUser):
    query = db.query(models.PortalFileReviewRecord).options(selectinload(models.PortalFileReviewRecord.created_by))
    if not current_user.is_superuser:
        query = query.filter(models.PortalFileReviewRecord.created_by_id == current_user.id)
    record = query.filter(models.PortalFileReviewRecord.id == record_id).first()
    if not record:
        raise HTTPException(status_code=404, detail="审查记录不存在")
    return serialize_portal_review_record(record, include_result=True)


@router.post("/reviews/rule-review/")
async def portal_rule_review(request: Request, db: DbSession, current_user: CurrentUser):
    payload = await request.json()
    document_id = str(payload.get("document_id") or "").strip()
    checklist_id = payload.get("checklist_id")
    if not document_id:
        raise HTTPException(status_code=400, detail="document_id不能为空")
    try:
        checklist_id = int(checklist_id)
    except (TypeError, ValueError):
        raise HTTPException(status_code=400, detail="checklist_id格式不正确")

    metadata = load_portal_metadata(document_id)
    checklist = portal_checklist_query(db).filter(
        models.FileReviewChecklist.id == checklist_id,
        models.FileReviewChecklist.is_deleted.is_(False),
    ).first()
    if not checklist:
        raise HTTPException(status_code=404, detail="审查清单不存在")

    checklist_payload = serialize_file_review_checklist(checklist)
    rules = [compact_rule_for_prompt(rule) for rule in checklist_payload.get("rules", [])]
    config = get_default_ai_config(db)
    ai_service = AIService(db, config=config)
    review_record = models.PortalFileReviewRecord(
        document_id=metadata["id"],
        document_name=metadata.get("original_name") or "",
        document_extension=metadata.get("extension") or "",
        document_size=int(metadata.get("size") or 0),
        page_count=int(metadata.get("page_count") or 0),
        word_count=int(metadata.get("word_count") or 0),
        checklist_id=checklist_payload["id"],
        checklist_name=checklist_payload["name"],
        rule_count=checklist_payload.get("rule_count") or len(rules),
        status="processing",
        model=ai_service.model,
        request_payload={
            "document_id": document_id,
            "checklist_id": checklist_id,
            "rule_count": len(rules),
        },
        created_by_id=current_user.id,
    )
    db.add(review_record)
    db.commit()
    db.refresh(review_record)

    if not rules:
        mark_portal_review_record_failed(db, review_record, "审查清单没有可用规则")
        raise HTTPException(status_code=400, detail="审查清单没有可用规则")
    if not ai_service.enabled:
        mark_portal_review_record_failed(db, review_record, "AI服务未启用或未配置API密钥")
        raise HTTPException(status_code=503, detail="AI服务未启用或未配置API密钥")

    page_results: list[dict[str, Any]] = []
    for page in metadata.get("pages", []):
        page_number = int(page.get("page_number") or len(page_results) + 1)
        chunks = split_review_chunks(page.get("text") or "")
        if not chunks:
            page_results.append(
                {
                    "page_number": page_number,
                    "chunk_count": 0,
                    "summary": "本页无可审查文本。",
                    "risk_level": "low",
                    "issues": [],
                    "raw_response": "",
                }
            )
            continue

        try:
            response_text = ai_service.call(
                build_rule_review_messages(metadata.get("original_name") or "", page, chunks, rules),
                max_tokens=3500,
                timeout=180,
            )
        except Exception as exc:
            error_message = f"大模型审查失败：第{page_number}页，{exc}"
            mark_portal_review_record_failed(db, review_record, error_message)
            raise HTTPException(status_code=502, detail=error_message) from exc
        page_results.append(normalize_review_page_result(page_number, chunks, response_text))

    issue_count = sum(len(page.get("issues") or []) for page in page_results)
    risk_order = {"low": 0, "medium": 1, "high": 2}
    overall_risk = "low"
    for page in page_results:
        risk = page.get("risk_level") or "low"
        if risk_order.get(risk, 0) > risk_order.get(overall_risk, 0):
            overall_risk = risk

    summary = {
        "page_count": len(page_results),
        "issue_count": issue_count,
        "risk_level": overall_risk,
    }
    review_record.status = "completed"
    review_record.summary = summary
    review_record.pages = page_results
    review_record.completed_at = datetime.now()
    db.add(review_record)
    db.commit()
    db.refresh(review_record)

    return {
        "success": True,
        "model": ai_service.model,
        "record": serialize_portal_review_record(review_record),
        "document": {
            "id": metadata["id"],
            "name": metadata["original_name"],
            "page_count": metadata.get("page_count") or len(page_results),
        },
        "checklist": {
            "id": checklist_payload["id"],
            "name": checklist_payload["name"],
            "rule_count": checklist_payload.get("rule_count") or len(rules),
        },
        "summary": summary,
        "pages": page_results,
    }


@router.post("/documents/upload/")
async def upload_portal_document(request: Request, file: UploadFile = File(...)):
    original_name = sanitize_filename(file.filename or "")
    extension = Path(original_name).suffix.lower()
    if extension not in PORTAL_ALLOWED_EXTENSIONS:
        raise HTTPException(status_code=400, detail="仅支持上传 DOCX、PDF 文件")

    document_id = uuid.uuid4().hex
    document_dir = portal_document_dir(document_id)
    document_dir.mkdir(parents=True, exist_ok=True)
    source_path = document_dir / f"source{extension}"

    size = 0
    try:
        with source_path.open("wb") as destination:
            while True:
                chunk = await file.read(1024 * 1024)
                if not chunk:
                    break
                size += len(chunk)
                if size > PORTAL_MAX_UPLOAD_SIZE:
                    raise HTTPException(status_code=400, detail="文件大小不能超过 20MB")
                destination.write(chunk)

        if size == 0:
            raise HTTPException(status_code=400, detail="不能上传空文件")

        validate_uploaded_file_signature(source_path, extension)
        preview = build_portal_preview(document_id, source_path, extension)
        content_type = (
            "application/pdf"
            if extension == ".pdf"
            else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        metadata = {
            "id": document_id,
            "original_name": original_name,
            "extension": extension[1:],
            "content_type": content_type,
            "size": size,
            "source_file": source_path.name,
            "title": preview.get("title") or original_name,
            "page_count": preview["page_count"],
            "word_count": preview["word_count"],
            "pages": preview["pages"],
            "created_at": datetime.now().isoformat(timespec="seconds"),
        }
        save_portal_metadata(metadata)
        return {"success": True, "message": "文件上传成功", "document": serialize_portal_document(request, metadata)}
    except HTTPException:
        shutil.rmtree(document_dir, ignore_errors=True)
        raise
    except Exception as exc:
        shutil.rmtree(document_dir, ignore_errors=True)
        raise HTTPException(status_code=500, detail=f"文件处理失败: {exc}") from exc
    finally:
        await file.close()


@router.get("/documents/{document_id}/preview/")
def get_portal_document_preview(document_id: str, request: Request):
    metadata = load_portal_metadata(document_id)
    return {"success": True, "document": serialize_portal_document(request, metadata)}


@router.get("/documents/{document_id}/file/")
def get_portal_document_file(document_id: str):
    metadata = load_portal_metadata(document_id)
    file_path = portal_document_dir(document_id) / metadata["source_file"]
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="源文件不存在")
    disposition = f"inline; filename*=UTF-8''{quote(metadata['original_name'])}"
    return FileResponse(
        file_path,
        media_type=metadata["content_type"],
        headers={"Content-Disposition": disposition},
    )


@router.get("/documents/{document_id}/pages/{page_number}.png")
def get_portal_pdf_page_image(document_id: str, page_number: int):
    metadata = load_portal_metadata(document_id)
    if metadata["extension"] != "pdf":
        raise HTTPException(status_code=404, detail="该文档没有 PDF 页图")
    page_count = int(metadata.get("page_count") or 0)
    if page_number < 1 or page_number > page_count:
        raise HTTPException(status_code=404, detail="页码不存在")
    if fitz is None:
        raise HTTPException(status_code=500, detail="服务端缺少 PyMuPDF 依赖")

    document_dir = portal_document_dir(document_id)
    pages_dir = document_dir / "pages"
    pages_dir.mkdir(parents=True, exist_ok=True)
    image_path = pages_dir / f"{page_number}.png"
    if not image_path.exists():
        source_path = document_dir / metadata["source_file"]
        if not source_path.exists():
            raise HTTPException(status_code=404, detail="源文件不存在")
        with fitz.open(source_path) as pdf:
            page = pdf[page_number - 1]
            pixmap = page.get_pixmap(matrix=fitz.Matrix(1.8, 1.8), alpha=False)
            pixmap.save(image_path)
    return FileResponse(image_path, media_type="image/png")
