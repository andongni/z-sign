import json
import logging
import re
import uuid
from datetime import datetime
from pathlib import Path
from typing import Any

import requests
from sqlalchemy.orm import Session

from app import models
from app.core.config import get_settings


logger = logging.getLogger(__name__)
settings = get_settings()


CONTRACT_TYPE_NAMES = {
    "procurement": "采购合同",
    "sales": "销售合同",
    "labor": "劳动合同",
    "service": "服务合同",
}


class AIService:
    def __init__(self, db: Session, config: models.AIModelConfig | None = None):
        if config is None:
            config = (
                db.query(models.AIModelConfig)
                .filter(models.AIModelConfig.is_default.is_(True), models.AIModelConfig.is_active.is_(True))
                .first()
            )
        self.config = config
        self.enabled = bool(config and config.is_active and config.api_key)
        self.api_key = config.api_key if config else ""
        self.api_url = f"{config.api_base_url.rstrip('/')}/chat/completions" if config else ""
        self.model = (
            config.default_model
            or ((config.available_models or [None])[0] if config else None)
            or ""
        )
        self.temperature = config.temperature if config else 0.7
        self.max_tokens = config.max_tokens if config else 2000
        self.timeout = config.timeout if config else 30

    def chat(self, message: str, history: list[dict[str, Any]] | None = None) -> str:
        messages = history or []
        messages = [item for item in messages if item.get("role") in {"system", "user", "assistant"}]
        messages.append({"role": "user", "content": message})
        return self.call(messages)

    def call(self, messages: list[dict[str, Any]], *, max_tokens: int | None = None, timeout: int | None = None) -> str:
        if not self.enabled:
            raise RuntimeError("AI服务未启用或未配置API密钥")
        if not self.model:
            raise RuntimeError("未设置默认模型")
        response = requests.post(
            self.api_url,
            headers={"Authorization": f"Bearer {self.api_key}", "Content-Type": "application/json"},
            json={
                "model": self.model,
                "messages": messages,
                "temperature": self.temperature,
                "max_tokens": max_tokens or self.max_tokens,
            },
            timeout=timeout or self.timeout,
        )
        if response.status_code != 200:
            raise RuntimeError(f"API调用失败（状态码：{response.status_code}）：{response.text[:300]}")
        data = response.json()
        try:
            return data["choices"][0]["message"]["content"]
        except (KeyError, IndexError, TypeError):
            raise RuntimeError("API响应格式错误")

    def test_connection(self) -> dict[str, Any]:
        content = self.call(
            [
                {"role": "system", "content": "你是一位专业的助手。"},
                {"role": "user", "content": "你好，请回复'连接成功'"},
            ],
            max_tokens=50,
            timeout=30,
        )
        return {"success": True, "message": "API连接测试成功", "response": content, "model": self.model}


def build_contract_generation_prompt(
    contract_type: str,
    industry: str = "",
    template: models.Template | None = None,
    basic_info: dict[str, Any] | None = None,
) -> str:
    basic_info = basic_info or {}
    contract_type_name = CONTRACT_TYPE_NAMES.get(contract_type, contract_type or "合同")
    parts = [
        "你是一位资深合同起草专家，请根据用户提供的信息生成一份规范、完整、可编辑的中文合同正文。",
        f"合同类型：{contract_type_name}",
    ]
    if industry:
        parts.append(f"所属行业：{industry}")
    if template:
        parts.append(f"参考模板：{template.name}\n{template.content}")
    if basic_info:
        parts.append("合同基本信息：")
        for key, val in basic_info.items():
            if val:
                parts.append(f"- {key}: {val}")
    parts.extend(
        [
            "要求：",
            "1. 使用正式合同语言，条款编号清晰。",
            "2. 包含合同主体、标的、金额/付款、履行、验收、违约责任、争议解决、签署信息等必要内容。",
            "3. 需要用户补充的信息用下划线占位。",
            "4. 直接输出合同正文，不要输出解释。",
        ]
    )
    return "\n".join(parts)


def generate_contract_content(
    db: Session,
    contract_type: str,
    industry: str = "",
    template: models.Template | None = None,
    basic_info: dict[str, Any] | None = None,
) -> dict[str, Any]:
    prompt = build_contract_generation_prompt(contract_type, industry, template, basic_info)
    try:
        text = AIService(db).call(
            [
                {"role": "system", "content": "你是一位专业合同起草专家。"},
                {"role": "user", "content": prompt},
            ],
            max_tokens=6000,
            timeout=240,
        )
    except Exception as exc:
        logger.warning("AI合同生成失败，返回基础模板: %s", exc)
        text = generate_fallback_contract(contract_type, industry, basic_info or {})
    return {"success": True, "content": {"text": text.strip()}, "message": "合同内容生成成功"}


def generate_fallback_contract(contract_type: str, industry: str, basic_info: dict[str, Any]) -> str:
    title = CONTRACT_TYPE_NAMES.get(contract_type, "合同")
    party_a = basic_info.get("party_a") or "________________"
    party_b = basic_info.get("party_b") or "________________"
    subject = basic_info.get("subject") or "________________"
    amount = basic_info.get("amount") or "________________"
    return f"""{title}

合同编号：________________

甲方：{party_a}
乙方：{party_b}

鉴于甲乙双方在平等、自愿、公平和诚实信用的基础上，就{subject}事项达成一致，订立本合同。

第一条 合同标的
1.1 本合同标的为：{subject}。
1.2 双方应根据本合同约定履行交付、验收、付款及配合义务。

第二条 合同金额及付款
2.1 合同金额为：{amount}。
2.2 付款方式、付款节点、发票要求由双方另行补充明确。

第三条 履行与验收
3.1 双方应按照约定时间、地点和质量标准履行合同。
3.2 验收中发现不符合约定的，责任方应在合理期限内完成整改。

第四条 违约责任
4.1 任一方违反本合同约定，应承担继续履行、采取补救措施或赔偿损失等责任。
4.2 因违约造成对方损失的，违约方应赔偿守约方因此遭受的实际损失。

第五条 争议解决
5.1 因本合同产生的争议，双方应先友好协商。
5.2 协商不成的，任一方可向有管辖权的人民法院提起诉讼。

第六条 其他
6.1 本合同未尽事宜，由双方签订补充协议。
6.2 本合同自双方签字或盖章之日起生效。

甲方（盖章）：________________
乙方（盖章）：________________

签订日期：________年____月____日
"""


def extract_contract_text(contract: models.Contract) -> str:
    if isinstance(contract.content, dict):
        return contract.content.get("text") or contract.content.get("html") or json.dumps(contract.content, ensure_ascii=False)
    if isinstance(contract.content, str):
        return contract.content
    return contract.title


def build_basic_review_result(contract: models.Contract) -> dict[str, Any]:
    text = extract_contract_text(contract)
    risk_items = []
    checks = [
        ("付款", "建议确认付款方式、付款节点和发票要求是否明确。", "medium"),
        ("违约", "建议核查违约责任是否包含计算方式和责任边界。", "medium"),
        ("争议", "建议明确争议解决方式和管辖机构。", "low"),
    ]
    for keyword, suggestion, level in checks:
        if keyword not in text:
            risk_items.append(
                {
                    "clause_id": keyword,
                    "clause_content": "",
                    "issue_description": f"未识别到明确的{keyword}相关条款",
                    "risk_level": level,
                    "legal_basis": "合同条款完整性要求",
                    "suggestion": suggestion,
                }
            )

    high_count = sum(1 for item in risk_items if item["risk_level"] == "high")
    medium_count = sum(1 for item in risk_items if item["risk_level"] == "medium")
    risk_level = "high" if high_count else ("medium" if medium_count else "low")
    score = max(60, 100 - high_count * 15 - medium_count * 8 - (len(risk_items) - high_count - medium_count) * 3)
    return {
        "overall_score": score,
        "risk_level": risk_level,
        "risk_count": len(risk_items),
        "summary": f"自动审核完成，发现{len(risk_items)}个需关注点",
        "review_data": {
            "contract_info": {
                "title": contract.title,
                "contract_no": contract.contract_no,
                "contract_type": CONTRACT_TYPE_NAMES.get(contract.contract_type, contract.contract_type),
                "industry": contract.industry,
            },
            "risk_overview": {
                "overall_score": score,
                "risk_level": risk_level,
                "risk_count": len(risk_items),
                "high_risk_count": high_count,
                "medium_risk_count": medium_count,
                "low_risk_count": len(risk_items) - high_count - medium_count,
            },
            "modification_suggestions": risk_items,
            "generated_at": datetime.now().isoformat(),
        },
        "issues": risk_items,
    }


def process_review_task(db: Session, task: models.ReviewTask) -> dict[str, Any]:
    contract = task.contract
    task.status = "ai_processing"
    task.started_at = datetime.now()
    task.progress = {
        "current_step": "自动审核",
        "progress": 50,
        "message": "正在生成审核结果",
        "steps": [{"name": "生成审核结果", "status": "processing"}],
    }
    db.flush()

    review_data = build_basic_review_result(contract)
    result = task.result
    if not result:
        result = models.ReviewResult(review_task_id=task.id, contract_id=contract.id)
        db.add(result)
        db.flush()

    result.overall_score = review_data["overall_score"]
    result.risk_level = review_data["risk_level"]
    result.risk_count = review_data["risk_count"]
    result.summary = review_data["summary"]
    result.review_data = review_data["review_data"]

    db.query(models.ReviewOpinion).filter(models.ReviewOpinion.review_result_id == result.id).delete()
    for issue in review_data["issues"]:
        db.add(
            models.ReviewOpinion(
                review_result_id=result.id,
                reviewer_id=task.reviewer_id,
                clause_id=issue["clause_id"],
                clause_content=issue["clause_content"],
                opinion_type="risk",
                risk_level=issue["risk_level"],
                opinion_content=issue["issue_description"],
                legal_basis=issue["legal_basis"],
                suggestion=issue["suggestion"],
                status="pending",
            )
        )

    if isinstance(task.review_levels, list) and task.review_levels:
        task.status = "manual_reviewing"
        task.completed_at = None
    else:
        task.status = "completed"
        task.completed_at = datetime.now()
    task.progress = {
        "current_step": "审核完成",
        "progress": 100,
        "message": "自动审核已完成",
        "steps": [{"name": "生成审核结果", "status": "completed"}],
    }
    db.commit()
    db.refresh(task)
    return {
        "success": True,
        "review_result_id": result.id,
        "status": task.status,
        "overall_score": review_data["overall_score"],
        "risk_level": review_data["risk_level"],
        "risk_count": review_data["risk_count"],
    }


def generate_reviewer_suggestions(
    db: Session,
    contract: models.Contract,
    reviewer: models.User,
    review_task: models.ReviewTask | None = None,
) -> dict[str, Any]:
    focus = (
        db.query(models.ReviewFocusConfig)
        .filter(models.ReviewFocusConfig.level == reviewer.reviewer_level, models.ReviewFocusConfig.is_active.is_(True))
        .first()
    )
    basic = build_basic_review_result(contract)
    suggestions = {
        "overall_evaluation": basic["summary"],
        "issues": basic["issues"],
        "focus_points": [
            {"point": item, "status": "需关注", "description": "请结合合同正文复核"}
            for item in ((focus.focus_points if focus else None) or [])
        ],
        "conclusion": "需要修改" if basic["issues"] else "通过",
        "summary": basic["summary"],
    }
    payload = {
        "reviewer_level": reviewer.reviewer_level,
        "reviewer_level_name": reviewer.reviewer_level or "",
        "focus_config": {
            "level": focus.level if focus else reviewer.reviewer_level,
            "level_name": focus.level_name if focus else "",
            "focus_points": focus.focus_points if focus else [],
            "review_standards": focus.review_standards if focus else "",
        },
        "suggestions": suggestions,
        "generated_at": datetime.now().isoformat(),
    }

    if review_task:
        result = review_task.result
        if not result:
            result = models.ReviewResult(
                review_task_id=review_task.id,
                contract_id=contract.id,
                overall_score=basic["overall_score"],
                risk_level=basic["risk_level"],
                risk_count=basic["risk_count"],
                summary=basic["summary"],
                review_data={},
            )
            db.add(result)
            db.flush()
        data = result.review_data or {}
        data["ai_suggestions"] = suggestions
        data["reviewer_level"] = reviewer.reviewer_level
        data["generated_at"] = payload["generated_at"]
        result.review_data = data
        db.commit()

    return payload


def write_review_report(result: models.ReviewResult, contract: models.Contract, report_format: str = "word") -> str:
    reports_dir = settings.media_root / "reports"
    reports_dir.mkdir(parents=True, exist_ok=True)
    suffix = ".docx" if report_format.lower() != "pdf" else ".txt"
    file_name = f"report_{result.id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:6]}{suffix}"
    output_path = reports_dir / file_name

    if suffix == ".docx":
        try:
            from docx import Document

            doc = Document()
            doc.add_heading("合同审核报告", level=1)
            doc.add_paragraph(f"合同名称：{contract.title}")
            doc.add_paragraph(f"合同编号：{contract.contract_no}")
            doc.add_paragraph(f"总体评分：{result.overall_score}")
            doc.add_paragraph(f"风险等级：{result.risk_level}")
            doc.add_paragraph(f"审核摘要：{result.summary}")
            doc.add_heading("审核意见", level=2)
            for opinion in result.opinions:
                doc.add_paragraph(f"[{opinion.risk_level}] {opinion.opinion_content} 建议：{opinion.suggestion}")
            doc.save(output_path)
        except Exception:
            output_path = output_path.with_suffix(".txt")
            output_path.write_text(build_report_text(result, contract), encoding="utf-8")
    else:
        output_path.write_text(build_report_text(result, contract), encoding="utf-8")

    return str(Path("reports") / output_path.name)


def build_report_text(result: models.ReviewResult, contract: models.Contract) -> str:
    lines = [
        "合同审核报告",
        f"合同名称：{contract.title}",
        f"合同编号：{contract.contract_no}",
        f"总体评分：{result.overall_score}",
        f"风险等级：{result.risk_level}",
        f"审核摘要：{result.summary}",
        "",
        "审核意见：",
    ]
    lines.extend(f"- [{opinion.risk_level}] {opinion.opinion_content} 建议：{opinion.suggestion}" for opinion in result.opinions)
    return "\n".join(lines)


def summarize_opinions(contract: models.Contract, tasks: list[models.ReviewTask]) -> dict[str, Any]:
    levels = {"level1": [], "level2": [], "level3": []}
    all_opinions = []
    for task in tasks:
        if not task.result:
            continue
        for opinion in task.result.opinions:
            item = {
                "id": opinion.id,
                "reviewer_level": task.reviewer_level,
                "type": opinion.opinion_type,
                "risk_level": opinion.risk_level,
                "content": opinion.opinion_content,
                "clause": opinion.clause_content,
                "suggestion": opinion.suggestion,
                "legal_basis": opinion.legal_basis,
                "status": opinion.status,
                "created_at": opinion.created_at.isoformat() if opinion.created_at else None,
            }
            all_opinions.append(item)
            if task.reviewer_level in levels:
                levels[task.reviewer_level].append(item)
    return {
        "contract_info": {
            "title": contract.title,
            "contract_no": contract.contract_no,
            "contract_type": CONTRACT_TYPE_NAMES.get(contract.contract_type, contract.contract_type),
        },
        "level1_opinions": levels["level1"],
        "level2_opinions": levels["level2"],
        "level3_opinions": levels["level3"],
        "all_opinions": all_opinions,
        "statistics": {
            "total_opinions": len(all_opinions),
            "high_risk_count": sum(1 for item in all_opinions if item["risk_level"] == "high"),
            "medium_risk_count": sum(1 for item in all_opinions if item["risk_level"] == "medium"),
            "low_risk_count": sum(1 for item in all_opinions if item["risk_level"] == "low"),
            "pending_count": sum(1 for item in all_opinions if item["status"] == "pending"),
        },
    }


def simple_recommendations_for_contract_type(contract_type: str) -> list[str]:
    return {
        "procurement": ["合同标的条款", "交货与验收条款", "付款方式条款", "质量保证条款", "违约责任条款"],
        "sales": ["合同标的条款", "交付条款", "验收条款", "付款条款", "质量保证条款"],
        "service": ["服务内容条款", "服务标准条款", "服务费用条款", "验收标准条款", "保密条款"],
        "labor": ["工作内容条款", "工作地点条款", "工作时间条款", "劳动报酬条款", "社会保险条款"],
    }.get(contract_type, ["合同主体条款", "履行条款", "付款条款", "违约责任条款", "争议解决条款"])


def split_text_to_clauses(text: str) -> list[str]:
    parts = re.split(r"(?:第[一二三四五六七八九十百]+条|\n\d+[.、])", text or "")
    return [part.strip() for part in parts if part.strip()]

