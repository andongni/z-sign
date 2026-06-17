from __future__ import annotations

import base64
import hashlib
import html
import hmac
import json
import os
import re
import secrets
import shutil
import time
import uuid
import zipfile
from dataclasses import dataclass
from datetime import date, datetime, timedelta
from decimal import Decimal
from pathlib import Path
from typing import Any, Callable
from urllib.parse import quote

import jwt
import MySQLdb
import requests
from dotenv import load_dotenv
from fastapi import Depends, FastAPI, File, HTTPException, Request, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, PlainTextResponse, StreamingResponse
from fastapi.security import HTTPAuthorizationCredentials, HTTPBearer
from MySQLdb.cursors import DictCursor

try:
    import docx
except ImportError:  # pragma: no cover - dependency is declared in requirements
    docx = None

try:
    import fitz
except ImportError:  # pragma: no cover - dependency is declared in requirements
    fitz = None

BASE_DIR = Path(__file__).resolve().parent.parent
load_dotenv(BASE_DIR / ".env")

SECRET_KEY = os.getenv("SECRET_KEY", "z_sign")
JWT_ALGORITHM = "HS256"
ACCESS_TOKEN_HOURS = int(os.getenv("ACCESS_TOKEN_HOURS", "2"))
REFRESH_TOKEN_DAYS = int(os.getenv("REFRESH_TOKEN_DAYS", "7"))
MEDIA_ROOT = BASE_DIR / "media"

app = FastAPI(title="Contract Review API", version="1.0.0")
security = HTTPBearer(auto_error=False)

cors_origins = os.getenv(
    "CORS_ALLOWED_ORIGINS",
    "http://localhost:5173,http://localhost:5174,http://localhost:3000,http://127.0.0.1:5173,http://127.0.0.1:5174,http://127.0.0.1:3000",
).split(",")

app.add_middleware(
    CORSMiddleware,
    allow_origins=[origin.strip() for origin in cors_origins if origin.strip()],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


def get_connection():
    return MySQLdb.connect(
        host=os.getenv("DB_HOST", "localhost"),
        port=int(os.getenv("DB_PORT", "3306")),
        user=os.getenv("DB_USER", "root"),
        passwd=os.getenv("DB_PASSWORD", ""),
        db=os.getenv("DB_NAME", "z_sign"),
        charset="utf8mb4",
        cursorclass=DictCursor,
        autocommit=False,
    )


def db_query(sql: str, params: tuple[Any, ...] = ()) -> list[dict[str, Any]]:
    conn = get_connection()
    try:
        with conn.cursor() as cursor:
            cursor.execute(sql, params)
            rows = cursor.fetchall()
        return [normalize_row(row) for row in rows]
    finally:
        conn.close()


def db_one(sql: str, params: tuple[Any, ...] = ()) -> dict[str, Any] | None:
    rows = db_query(sql, params)
    return rows[0] if rows else None


def db_execute(sql: str, params: tuple[Any, ...] = ()) -> int:
    conn = get_connection()
    try:
        with conn.cursor() as cursor:
            cursor.execute(sql, params)
            last_id = cursor.lastrowid
        conn.commit()
        return int(last_id or 0)
    except Exception:
        conn.rollback()
        raise
    finally:
        conn.close()


def normalize_value(value: Any) -> Any:
    if isinstance(value, Decimal):
        return float(value)
    if isinstance(value, (datetime, date)):
        return value.isoformat()
    if isinstance(value, bytes):
        return value.decode("utf-8", errors="ignore")
    return value


def normalize_row(row: dict[str, Any]) -> dict[str, Any]:
    normalized = {key: normalize_value(value) for key, value in row.items()}
    for key, value in list(normalized.items()):
        if isinstance(value, str) and key in JSON_FIELD_NAMES:
            try:
                normalized[key] = json.loads(value) if value else None
            except json.JSONDecodeError:
                pass
    return normalized


JSON_FIELD_NAMES = {
    "content",
    "tags",
    "habit_value",
    "request_data",
    "response_data",
    "review_levels",
    "reviewer_assignments",
    "progress",
    "review_data",
    "focus_points",
    "attention_items",
    "available_models",
    "rule_content",
    "match_result",
    "properties",
    "relation_properties",
    "related_clauses",
    "extracted_data",
    "result_data",
}


def encode_db_value(key: str, value: Any) -> Any:
    if key in JSON_FIELD_NAMES and isinstance(value, (dict, list)):
        return json.dumps(value, ensure_ascii=False)
    if isinstance(value, bool):
        return int(value)
    return value


def now() -> str:
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")


def generate_review_rule_code() -> str:
    date_part = datetime.now().strftime("%Y%m%d")
    timestamp_tail = int(time.time() * 1000) % 100000
    for offset in range(100000):
        code = f"{date_part}{(timestamp_tail + offset) % 100000:05d}"
        if not db_one("SELECT id FROM rules_review_rule WHERE rule_code = %s LIMIT 1", (code,)):
            return code
    raise HTTPException(status_code=400, detail="规则编码生成失败，请稍后重试")


def make_token(user: dict[str, Any], token_type: str) -> str:
    if token_type == "refresh":
        expires = datetime.utcnow() + timedelta(days=REFRESH_TOKEN_DAYS)
    else:
        expires = datetime.utcnow() + timedelta(hours=ACCESS_TOKEN_HOURS)
    payload = {
        "type": token_type,
        "user_id": user["id"],
        "username": user["username"],
        "exp": expires,
        "iat": datetime.utcnow(),
    }
    return jwt.encode(payload, SECRET_KEY, algorithm=JWT_ALGORITHM)


def decode_token(token: str) -> dict[str, Any]:
    try:
        return jwt.decode(token, SECRET_KEY, algorithms=[JWT_ALGORITHM])
    except jwt.PyJWTError as exc:
        raise HTTPException(status_code=401, detail="Invalid token") from exc


def verify_django_password(password: str, encoded: str) -> bool:
    try:
        algorithm, iterations, salt, expected = encoded.split("$", 3)
    except ValueError:
        return False
    if algorithm != "pbkdf2_sha256":
        return False
    derived = hashlib.pbkdf2_hmac(
        "sha256",
        password.encode(),
        salt.encode(),
        int(iterations),
    )
    actual = base64.b64encode(derived).decode().strip()
    return hmac.compare_digest(actual, expected)


def make_django_password(password: str) -> str:
    salt = secrets.token_urlsafe(12)[:12]
    iterations = 720000
    derived = hashlib.pbkdf2_hmac("sha256", password.encode(), salt.encode(), iterations)
    digest = base64.b64encode(derived).decode().strip()
    return f"pbkdf2_sha256${iterations}${salt}${digest}"


def current_user(
    credentials: HTTPAuthorizationCredentials | None = Depends(security),
) -> dict[str, Any]:
    if not credentials:
        raise HTTPException(status_code=401, detail="Authentication credentials were not provided.")
    payload = decode_token(credentials.credentials)
    user = get_user_by_id(payload.get("user_id"))
    if not user or not user.get("is_active") or user.get("is_deleted"):
        raise HTTPException(status_code=401, detail="User is inactive or deleted.")
    return serialize_user(user)


def get_user_by_id(user_id: int | str | None) -> dict[str, Any] | None:
    if not user_id:
        return None
    return db_one(
        """
        SELECT u.*, d.name AS department_name
        FROM users_user u
        LEFT JOIN users_department d ON d.id = u.department_id
        WHERE u.id = %s
        """,
        (user_id,),
    )


def get_roles_for_user(user_id: int) -> list[dict[str, Any]]:
    return db_query(
        """
        SELECT r.id, r.name, r.code, r.description
        FROM users_user_role ur
        JOIN users_role r ON r.id = ur.role_id
        WHERE ur.user_id = %s AND COALESCE(r.is_deleted, 0) = 0
        ORDER BY r.id
        """,
        (user_id,),
    )


def serialize_user(row: dict[str, Any]) -> dict[str, Any]:
    data = public_fields(
        row,
        [
            "id",
            "username",
            "email",
            "real_name",
            "phone",
            "avatar",
            "department_id",
            "department_name",
            "role",
            "reviewer_level",
            "is_active",
            "created_at",
            "updated_at",
        ],
    )
    data["department"] = data.pop("department_id", None)
    data["roles"] = get_roles_for_user(int(row["id"]))
    return data


def public_fields(row: dict[str, Any], fields: list[str]) -> dict[str, Any]:
    return {field: row.get(field) for field in fields if field in row}


def rows_by_ids(table: str, ids: list[Any]) -> dict[int, dict[str, Any]]:
    clean_ids = [int(item) for item in ids if item is not None]
    if not clean_ids:
        return {}
    placeholders = ", ".join(["%s"] * len(clean_ids))
    rows = db_query(f"SELECT * FROM {table} WHERE id IN ({placeholders})", tuple(clean_ids))
    return {int(row["id"]): row for row in rows}


def filter_payload(payload: dict[str, Any], allowed: list[str], aliases: dict[str, str] | None = None):
    aliases = aliases or {}
    data: dict[str, Any] = {}
    for key, value in payload.items():
        mapped = aliases.get(key, key)
        if mapped in allowed:
            data[mapped] = encode_db_value(mapped, value)
    return data


def insert_row(table: str, data: dict[str, Any]) -> dict[str, Any]:
    columns = list(data.keys())
    placeholders = ", ".join(["%s"] * len(columns))
    column_sql = ", ".join(columns)
    last_id = db_execute(
        f"INSERT INTO {table} ({column_sql}) VALUES ({placeholders})",
        tuple(data[col] for col in columns),
    )
    row = db_one(f"SELECT * FROM {table} WHERE id = %s", (last_id,))
    if not row:
        raise HTTPException(status_code=500, detail="Insert failed")
    return row


def update_row(table: str, row_id: int, data: dict[str, Any]) -> dict[str, Any]:
    if not data:
        row = db_one(f"SELECT * FROM {table} WHERE id = %s", (row_id,))
        if not row:
            raise HTTPException(status_code=404, detail="Not found")
        return row
    assignments = ", ".join([f"{key} = %s" for key in data])
    params = tuple(data.values()) + (row_id,)
    db_execute(f"UPDATE {table} SET {assignments} WHERE id = %s", params)
    row = db_one(f"SELECT * FROM {table} WHERE id = %s", (row_id,))
    if not row:
        raise HTTPException(status_code=404, detail="Not found")
    return row


def soft_delete_or_delete(table: str, row_id: int, soft: bool = True) -> dict[str, str]:
    if soft:
        db_execute(f"UPDATE {table} SET is_deleted = 1 WHERE id = %s", (row_id,))
    else:
        db_execute(f"DELETE FROM {table} WHERE id = %s", (row_id,))
    return {"message": "Deleted"}


@dataclass(frozen=True)
class ResourceSpec:
    table: str
    fields: list[str]
    create_fields: list[str]
    update_fields: list[str]
    aliases: dict[str, str]
    filters: list[str]
    search: list[str]
    default_order: str
    soft_delete: bool = True
    serializer: Callable[[dict[str, Any]], dict[str, Any]] | None = None


def list_resource(request: Request, spec: ResourceSpec) -> dict[str, Any]:
    query = request.query_params
    where: list[str] = []
    params: list[Any] = []
    if spec.soft_delete and "is_deleted" in spec.fields:
        where.append("COALESCE(is_deleted, 0) = 0")
    for field in spec.filters:
        value = query.get(field)
        column = spec.aliases.get(field, field)
        if value not in (None, ""):
            where.append(f"{column} = %s")
            params.append(parse_query_value(value))
    search = query.get("search", "").strip()
    if search and spec.search:
        parts = [f"{field} LIKE %s" for field in spec.search]
        where.append("(" + " OR ".join(parts) + ")")
        params.extend([f"%{search}%"] * len(spec.search))
    where_sql = " WHERE " + " AND ".join(where) if where else ""
    count_row = db_one(f"SELECT COUNT(*) AS count FROM {spec.table}{where_sql}", tuple(params))
    count = int(count_row["count"] if count_row else 0)
    page = max(int(query.get("page", "1") or 1), 1)
    page_size = min(max(int(query.get("page_size", "20") or 20), 1), 1000)
    offset = (page - 1) * page_size
    ordering = query.get("ordering") or spec.default_order
    order_sql = build_order_sql(ordering, spec.fields)
    rows = db_query(
        f"SELECT * FROM {spec.table}{where_sql}{order_sql} LIMIT %s OFFSET %s",
        tuple(params) + (page_size, offset),
    )
    serializer = spec.serializer or (lambda row: row)
    return {
        "count": count,
        "next": None,
        "previous": None,
        "results": [serializer(row) for row in rows],
    }


def parse_query_value(value: str) -> Any:
    if value.lower() == "true":
        return 1
    if value.lower() == "false":
        return 0
    return value


def build_order_sql(ordering: str, allowed_fields: list[str]) -> str:
    if not ordering:
        return ""
    direction = "ASC"
    field = ordering
    if ordering.startswith("-"):
        direction = "DESC"
        field = ordering[1:]
    if field not in allowed_fields:
        return ""
    return f" ORDER BY {field} {direction}"


def get_resource(spec: ResourceSpec, row_id: int) -> dict[str, Any]:
    row = db_one(f"SELECT * FROM {spec.table} WHERE id = %s", (row_id,))
    if not row or (spec.soft_delete and row.get("is_deleted")):
        raise HTTPException(status_code=404, detail="Not found")
    serializer = spec.serializer or (lambda item: item)
    return serializer(row)


def create_resource(spec: ResourceSpec, payload: dict[str, Any], extra: dict[str, Any] | None = None):
    data = filter_payload(payload, spec.create_fields, spec.aliases)
    if spec.table == "rules_review_rule":
        data["rule_code"] = generate_review_rule_code()
    data.update(extra or {})
    if "created_at" in spec.fields:
        data.setdefault("created_at", now())
    if "updated_at" in spec.fields:
        data.setdefault("updated_at", now())
    row = insert_row(spec.table, data)
    return (spec.serializer or (lambda item: item))(row)


def update_resource(spec: ResourceSpec, row_id: int, payload: dict[str, Any]):
    data = filter_payload(payload, spec.update_fields, spec.aliases)
    if "updated_at" in spec.fields:
        data["updated_at"] = now()
    row = update_row(spec.table, row_id, data)
    return (spec.serializer or (lambda item: item))(row)


def serialize_contract(row: dict[str, Any]) -> dict[str, Any]:
    data = public_fields(
        row,
        [
            "id",
            "contract_no",
            "title",
            "contract_type",
            "industry",
            "status",
            "content",
            "file_path",
            "file_format",
            "template_id",
            "drafter_id",
            "current_version",
            "created_at",
            "updated_at",
        ],
    )
    data["template"] = data.pop("template_id", None)
    data["drafter"] = data.pop("drafter_id", None)
    data["drafter_name"] = row.get("drafter_name") or lookup_name("users_user", data.get("drafter"))
    data["template_name"] = row.get("template_name") or lookup_name("contracts_template", data.get("template"), "name")
    data["versions"] = db_query(
        "SELECT * FROM contracts_contract_version WHERE contract_id = %s AND COALESCE(is_deleted, 0) = 0 ORDER BY version DESC",
        (row["id"],),
    )
    return data


def lookup_name(table: str, row_id: Any, field: str = "username") -> str:
    if not row_id:
        return ""
    row = db_one(f"SELECT {field} FROM {table} WHERE id = %s", (row_id,))
    return str(row.get(field) or "") if row else ""


def serialize_template(row: dict[str, Any]) -> dict[str, Any]:
    data = public_fields(
        row,
        [
            "id",
            "name",
            "contract_type",
            "industry",
            "category",
            "content",
            "description",
            "tags",
            "usage_count",
            "is_public",
            "is_enterprise",
            "created_by_id",
            "enterprise_id",
            "created_at",
            "updated_at",
        ],
    )
    data["created_by"] = data.pop("created_by_id", None)
    data["created_by_name"] = lookup_name("users_user", data.get("created_by"))
    return data


def serialize_review_result(row: dict[str, Any]) -> dict[str, Any]:
    opinions = db_query(
        "SELECT * FROM reviews_review_opinion WHERE review_result_id = %s AND COALESCE(is_deleted, 0) = 0 ORDER BY created_at DESC",
        (row["id"],),
    )
    data = public_fields(
        row,
        [
            "id",
            "review_task_id",
            "contract_id",
            "overall_score",
            "risk_level",
            "risk_count",
            "summary",
            "report_path",
            "report_format",
            "review_data",
            "created_at",
        ],
    )
    data["review_task"] = data.pop("review_task_id", None)
    data["contract"] = data.pop("contract_id", None)
    data["contract_title"] = lookup_name("contracts_contract", data.get("contract"), "title")
    data["opinions"] = [serialize_review_opinion(item) for item in opinions]
    return data


def serialize_review_opinion(row: dict[str, Any]) -> dict[str, Any]:
    data = public_fields(
        row,
        [
            "id",
            "review_result_id",
            "reviewer_id",
            "clause_id",
            "clause_content",
            "opinion_type",
            "risk_level",
            "opinion_content",
            "legal_basis",
            "suggestion",
            "status",
            "created_at",
            "updated_at",
        ],
    )
    data["review_result"] = data.pop("review_result_id", None)
    data["reviewer"] = data.pop("reviewer_id", None)
    data["reviewer_name"] = lookup_name("users_user", data.get("reviewer"))
    return data


def serialize_review_task(row: dict[str, Any]) -> dict[str, Any]:
    result = db_one("SELECT * FROM reviews_review_result WHERE review_task_id = %s", (row["id"],))
    data = public_fields(
        row,
        [
            "id",
            "contract_id",
            "contract_version",
            "status",
            "priority",
            "reviewer_id",
            "reviewer_level",
            "review_levels",
            "reviewer_assignments",
            "celery_task_id",
            "progress",
            "started_at",
            "completed_at",
            "error_message",
            "created_by_id",
            "created_at",
            "updated_at",
        ],
    )
    data["contract"] = data.pop("contract_id", None)
    data["reviewer"] = data.pop("reviewer_id", None)
    data["created_by"] = data.pop("created_by_id", None)
    data["contract_title"] = lookup_name("contracts_contract", data.get("contract"), "title")
    data["reviewer_name"] = lookup_name("users_user", data.get("reviewer"))
    data["created_by_name"] = lookup_name("users_user", data.get("created_by"))
    data["reviewer_assignments_detail"] = reviewer_assignment_detail(data.get("reviewer_assignments"))
    data["result"] = serialize_review_result(result) if result else None
    return data


def reviewer_assignment_detail(assignments: Any) -> dict[str, Any]:
    if not isinstance(assignments, dict):
        return {}
    details: dict[str, Any] = {}
    for level, user_id in assignments.items():
        user = get_user_by_id(user_id)
        details[level] = {
            "id": user.get("id") if user else user_id,
            "username": user.get("username") if user else "unknown",
            "real_name": (user.get("real_name") or user.get("username")) if user else "unknown",
            "email": user.get("email") if user else "",
        }
    return details


def serialize_ai_config(row: dict[str, Any]) -> dict[str, Any]:
    data = row.copy()
    data["created_by"] = data.pop("created_by_id", None)
    data["updated_by"] = data.pop("updated_by_id", None)
    data["created_by_name"] = lookup_name("users_user", data.get("created_by"))
    data["updated_by_name"] = lookup_name("users_user", data.get("updated_by"))
    data["provider_display"] = data.get("provider") or ""
    return data


def serialize_relation(row: dict[str, Any]) -> dict[str, Any]:
    data = row.copy()
    data["source_entity"] = data.pop("source_entity_id", None)
    data["target_entity"] = data.pop("target_entity_id", None)
    data["source_entity_id"] = data["source_entity"]
    data["target_entity_id"] = data["target_entity"]
    data["source_entity_name"] = lookup_name("knowledge_knowledge_entity", data["source_entity"], "entity_name")
    data["target_entity_name"] = lookup_name("knowledge_knowledge_entity", data["target_entity"], "entity_name")
    return data


def serialize_role(row: dict[str, Any]) -> dict[str, Any]:
    permissions = db_query(
        """
        SELECT p.*
        FROM users_role_permission rp
        JOIN users_permission p ON p.id = rp.permission_id
        WHERE rp.role_id = %s AND COALESCE(p.is_deleted, 0) = 0
        """,
        (row["id"],),
    )
    data = row.copy()
    data["permissions"] = permissions
    return data


def serialize_clause(row: dict[str, Any]) -> dict[str, Any]:
    data = row.copy()
    data["contract"] = data.pop("contract_id", None)
    data["confirmed_by"] = data.pop("confirmed_by_id", None)
    data["contract_title"] = lookup_name("contracts_contract", data.get("contract"), "title")
    data["confirmed_by_name"] = lookup_name("users_user", data.get("confirmed_by"))
    return data


def serialize_risk(row: dict[str, Any]) -> dict[str, Any]:
    data = row.copy()
    data["review_result"] = data.pop("review_result_id", None)
    data["clause"] = data.pop("clause_id", None)
    data["handled_by"] = data.pop("handled_by_id", None)
    data["handled_by_name"] = lookup_name("users_user", data.get("handled_by"))
    data["clause_content"] = lookup_name("clauses_contract_clause", data.get("clause"), "clause_content")
    return data


def serialize_comparison_diff(row: dict[str, Any]) -> dict[str, Any]:
    data = row.copy()
    data["comparison_task"] = data.pop("comparison_task_id", None)
    return data


def serialize_comparison_task(row: dict[str, Any]) -> dict[str, Any]:
    diffs = db_query(
        "SELECT * FROM comparisons_comparison_diff WHERE comparison_task_id = %s ORDER BY risk_level DESC, created_at ASC",
        (row["id"],),
    )
    data = row.copy()
    data["source_contract"] = data.pop("source_contract_id", None)
    data["target_contract"] = data.pop("target_contract_id", None)
    data["template"] = data.pop("template_id", None)
    data["created_by"] = data.pop("created_by_id", None)
    data["source_contract_title"] = lookup_name("contracts_contract", data.get("source_contract"), "title")
    data["target_contract_title"] = lookup_name("contracts_contract", data.get("target_contract"), "title")
    data["template_name"] = lookup_name("contracts_template", data.get("template"), "name")
    data["created_by_name"] = lookup_name("users_user", data.get("created_by"))
    data["diffs"] = [serialize_comparison_diff(diff) for diff in diffs]
    return data


def serialize_recommendation(row: dict[str, Any]) -> dict[str, Any]:
    data = row.copy()
    data["user"] = data.pop("user_id", None)
    data["contract"] = data.pop("contract_id", None)
    data["user_name"] = lookup_name("users_user", data.get("user"))
    data["contract_title"] = lookup_name("contracts_contract", data.get("contract"), "title")
    return data


USER_FIELDS = [
    "id",
    "password",
    "last_login",
    "is_superuser",
    "username",
    "email",
    "real_name",
    "phone",
    "avatar",
    "role",
    "is_active",
    "is_staff",
    "is_deleted",
    "created_at",
    "updated_at",
    "department_id",
    "reviewer_level",
]

SPECS = {
    "departments": ResourceSpec(
        "users_department",
        ["id", "name", "parent_id", "code", "description", "is_deleted", "created_at", "updated_at"],
        ["name", "parent_id", "code", "description"],
        ["name", "parent_id", "code", "description"],
        {"parent": "parent_id"},
        ["parent", "code"],
        ["name", "code"],
        "-created_at",
    ),
    "roles": ResourceSpec(
        "users_role",
        ["id", "name", "code", "description", "is_deleted", "created_at", "updated_at"],
        ["name", "code", "description"],
        ["name", "code", "description"],
        {},
        ["code"],
        ["name", "code"],
        "-created_at",
        serializer=serialize_role,
    ),
    "permissions": ResourceSpec(
        "users_permission",
        ["id", "name", "code", "resource", "action", "description", "is_deleted", "created_at", "updated_at"],
        ["name", "code", "resource", "action", "description"],
        ["name", "code", "resource", "action", "description"],
        {},
        ["resource", "action"],
        ["name", "code"],
        "-created_at",
    ),
    "contracts": ResourceSpec(
        "contracts_contract",
        [
            "id",
            "contract_no",
            "title",
            "contract_type",
            "industry",
            "status",
            "content",
            "file_path",
            "file_format",
            "template_id",
            "drafter_id",
            "current_version",
            "is_deleted",
            "created_at",
            "updated_at",
        ],
        ["title", "contract_type", "industry", "content", "file_path", "file_format", "template_id", "status"],
        ["title", "contract_type", "industry", "content", "file_path", "file_format", "template_id", "status"],
        {"template": "template_id", "drafter": "drafter_id"},
        ["contract_type", "industry", "status", "drafter"],
        ["title", "contract_no"],
        "-created_at",
        serializer=serialize_contract,
    ),
    "templates": ResourceSpec(
        "contracts_template",
        [
            "id",
            "name",
            "contract_type",
            "industry",
            "category",
            "content",
            "description",
            "tags",
            "usage_count",
            "is_public",
            "is_enterprise",
            "created_by_id",
            "enterprise_id",
            "is_deleted",
            "created_at",
            "updated_at",
        ],
        ["name", "contract_type", "industry", "category", "content", "description", "tags", "is_public", "is_enterprise", "enterprise_id"],
        ["name", "contract_type", "industry", "category", "content", "description", "tags", "is_public", "is_enterprise", "enterprise_id"],
        {"created_by": "created_by_id"},
        ["contract_type", "industry", "is_public", "is_enterprise"],
        ["name", "description"],
        "-usage_count",
        serializer=serialize_template,
    ),
    "review_tasks": ResourceSpec(
        "reviews_review_task",
        [
            "id",
            "contract_id",
            "contract_version",
            "task_type",
            "status",
            "priority",
            "reviewer_id",
            "reviewer_level",
            "review_levels",
            "reviewer_assignments",
            "celery_task_id",
            "progress",
            "started_at",
            "completed_at",
            "error_message",
            "created_by_id",
            "created_at",
            "updated_at",
        ],
        ["contract_id", "contract_version", "task_type", "status", "priority", "reviewer_id", "reviewer_level", "review_levels", "reviewer_assignments"],
        ["contract_id", "contract_version", "task_type", "status", "priority", "reviewer_id", "reviewer_level", "review_levels", "reviewer_assignments", "progress", "error_message"],
        {"contract": "contract_id", "reviewer": "reviewer_id", "created_by": "created_by_id"},
        ["contract", "status"],
        [],
        "-created_at",
        soft_delete=False,
        serializer=serialize_review_task,
    ),
    "review_focus_configs": ResourceSpec(
        "reviews_review_focus_config",
        [
            "id",
            "level",
            "level_name",
            "focus_points",
            "focus_description",
            "review_standards",
            "attention_items",
            "is_active",
            "created_by_id",
            "updated_by_id",
            "created_at",
            "updated_at",
        ],
        ["level", "level_name", "focus_points", "focus_description", "review_standards", "attention_items", "is_active"],
        ["level", "level_name", "focus_points", "focus_description", "review_standards", "attention_items", "is_active"],
        {"created_by": "created_by_id", "updated_by": "updated_by_id"},
        ["level", "is_active"],
        ["level_name", "focus_description"],
        "level",
        soft_delete=False,
    ),
    "ai_model_configs": ResourceSpec(
        "reviews_ai_model_config",
        [
            "id",
            "name",
            "provider",
            "api_key",
            "api_base_url",
            "available_models",
            "default_model",
            "is_active",
            "is_default",
            "description",
            "temperature",
            "max_tokens",
            "timeout",
            "created_by_id",
            "updated_by_id",
            "created_at",
            "updated_at",
        ],
        ["name", "provider", "api_key", "api_base_url", "available_models", "default_model", "is_active", "is_default", "description", "temperature", "max_tokens", "timeout"],
        ["name", "provider", "api_key", "api_base_url", "available_models", "default_model", "is_active", "is_default", "description", "temperature", "max_tokens", "timeout"],
        {"created_by": "created_by_id", "updated_by": "updated_by_id"},
        ["provider", "is_active", "is_default"],
        ["name", "description"],
        "-created_at",
        soft_delete=False,
        serializer=serialize_ai_config,
    ),
    "rules": ResourceSpec(
        "rules_review_rule",
        [
            "id",
            "rule_code",
            "rule_name",
            "rule_type",
            "industry",
            "category",
            "priority",
            "rule_content",
            "risk_level",
            "legal_basis",
            "description",
            "is_active",
            "version",
            "created_by_id",
            "is_deleted",
            "created_at",
            "updated_at",
        ],
        ["rule_code", "rule_name", "rule_type", "industry", "category", "priority", "rule_content", "risk_level", "legal_basis", "description", "is_active", "version"],
        ["rule_name", "rule_type", "industry", "category", "priority", "rule_content", "risk_level", "legal_basis", "description", "is_active", "version"],
        {"created_by": "created_by_id"},
        ["rule_type", "risk_level", "is_active"],
        ["rule_code", "rule_name", "description"],
        "-priority",
    ),
    "matches": ResourceSpec(
        "rules_rule_match",
        ["id", "review_task_id", "rule_id", "contract_id", "matched_clause", "match_score", "match_result", "created_at"],
        ["review_task_id", "rule_id", "contract_id", "matched_clause", "match_score", "match_result"],
        ["review_task_id", "rule_id", "contract_id", "matched_clause", "match_score", "match_result"],
        {"review_task": "review_task_id", "rule": "rule_id"},
        ["review_task", "rule", "contract_id"],
        [],
        "-match_score",
        soft_delete=False,
    ),
    "entities": ResourceSpec(
        "knowledge_knowledge_entity",
        ["id", "entity_type", "entity_name", "entity_code", "description", "properties", "source", "is_deleted", "created_at", "updated_at"],
        ["entity_type", "entity_name", "entity_code", "description", "properties", "source"],
        ["entity_type", "entity_name", "entity_code", "description", "properties", "source"],
        {},
        ["entity_type"],
        ["entity_name", "entity_code", "description"],
        "entity_type",
    ),
    "relations": ResourceSpec(
        "knowledge_knowledge_relation",
        ["id", "source_entity_id", "target_entity_id", "relation_type", "relation_properties", "confidence", "is_deleted", "created_at", "updated_at"],
        ["source_entity_id", "target_entity_id", "relation_type", "relation_properties", "confidence"],
        ["source_entity_id", "target_entity_id", "relation_type", "relation_properties", "confidence"],
        {"source_entity": "source_entity_id", "target_entity": "target_entity_id"},
        ["relation_type"],
        [],
        "-created_at",
        serializer=serialize_relation,
    ),
    "regulations": ResourceSpec(
        "knowledge_regulation",
        ["id", "title", "regulation_no", "regulation_type", "publish_date", "effective_date", "expiry_date", "content", "source_url", "entity_id", "is_active", "is_deleted", "created_at", "updated_at"],
        ["title", "regulation_no", "regulation_type", "publish_date", "effective_date", "expiry_date", "content", "source_url", "entity_id", "is_active"],
        ["title", "regulation_no", "regulation_type", "publish_date", "effective_date", "expiry_date", "content", "source_url", "entity_id", "is_active"],
        {"entity": "entity_id"},
        ["regulation_type", "is_active"],
        ["title", "regulation_no", "content"],
        "-publish_date",
    ),
    "cases": ResourceSpec(
        "knowledge_case",
        ["id", "case_no", "case_title", "case_type", "court", "judge_date", "case_summary", "case_content", "related_clauses", "entity_id", "is_deleted", "created_at", "updated_at"],
        ["case_no", "case_title", "case_type", "court", "judge_date", "case_summary", "case_content", "related_clauses", "entity_id"],
        ["case_no", "case_title", "case_type", "court", "judge_date", "case_summary", "case_content", "related_clauses", "entity_id"],
        {"entity": "entity_id"},
        ["case_type"],
        ["case_no", "case_title", "court"],
        "-judge_date",
    ),
    "clauses": ResourceSpec(
        "clauses_contract_clause",
        ["id", "contract_id", "contract_version", "clause_no", "clause_type", "clause_title", "clause_content", "start_position", "end_position", "extracted_data", "confidence", "is_confirmed", "confirmed_by_id", "confirmed_at", "created_at", "updated_at"],
        ["contract_id", "contract_version", "clause_no", "clause_type", "clause_title", "clause_content", "start_position", "end_position", "extracted_data", "confidence", "is_confirmed"],
        ["contract_id", "contract_version", "clause_no", "clause_type", "clause_title", "clause_content", "start_position", "end_position", "extracted_data", "confidence", "is_confirmed"],
        {"contract": "contract_id", "confirmed_by": "confirmed_by_id"},
        ["contract", "contract_version", "clause_type", "is_confirmed"],
        ["clause_no", "clause_title", "clause_content"],
        "start_position",
        soft_delete=False,
        serializer=serialize_clause,
    ),
    "risks": ResourceSpec(
        "risks_risk_identification",
        ["id", "review_result_id", "contract_id", "clause_id", "risk_type", "risk_category", "risk_level", "risk_description", "risk_location", "legal_basis", "suggestion", "status", "handled_by_id", "handled_at", "created_at"],
        ["review_result_id", "contract_id", "clause_id", "risk_type", "risk_category", "risk_level", "risk_description", "risk_location", "legal_basis", "suggestion", "status"],
        ["review_result_id", "contract_id", "clause_id", "risk_type", "risk_category", "risk_level", "risk_description", "risk_location", "legal_basis", "suggestion", "status"],
        {"review_result": "review_result_id", "clause": "clause_id", "handled_by": "handled_by_id"},
        ["review_result", "contract_id", "risk_type", "risk_category", "risk_level", "status"],
        ["risk_description", "risk_location", "legal_basis", "suggestion"],
        "-risk_level",
        soft_delete=False,
        serializer=serialize_risk,
    ),
    "comparison_tasks": ResourceSpec(
        "comparisons_comparison_task",
        ["id", "task_type", "source_contract_id", "target_contract_id", "source_version", "target_version", "template_id", "status", "result_data", "created_by_id", "created_at", "completed_at"],
        ["task_type", "source_contract_id", "target_contract_id", "source_version", "target_version", "template_id", "status", "result_data"],
        ["task_type", "source_contract_id", "target_contract_id", "source_version", "target_version", "template_id", "status", "result_data", "completed_at"],
        {"source_contract": "source_contract_id", "target_contract": "target_contract_id", "template": "template_id", "created_by": "created_by_id"},
        ["task_type", "status"],
        [],
        "-created_at",
        soft_delete=False,
        serializer=serialize_comparison_task,
    ),
    "comparison_diffs": ResourceSpec(
        "comparisons_comparison_diff",
        ["id", "comparison_task_id", "diff_type", "diff_level", "source_content", "target_content", "clause_id", "risk_level", "created_at"],
        ["comparison_task_id", "diff_type", "diff_level", "source_content", "target_content", "clause_id", "risk_level"],
        ["comparison_task_id", "diff_type", "diff_level", "source_content", "target_content", "clause_id", "risk_level"],
        {"comparison_task": "comparison_task_id"},
        ["comparison_task", "diff_type", "diff_level", "risk_level"],
        ["source_content", "target_content", "clause_id"],
        "-risk_level",
        soft_delete=False,
        serializer=serialize_comparison_diff,
    ),
    "recommendations": ResourceSpec(
        "recommendations_recommendation",
        ["id", "user_id", "contract_id", "recommendation_type", "recommendation_context", "item_type", "item_id", "item_content", "score", "reason", "is_accepted", "created_at"],
        ["user_id", "contract_id", "recommendation_type", "recommendation_context", "item_type", "item_id", "item_content", "score", "reason", "is_accepted"],
        ["user_id", "contract_id", "recommendation_type", "recommendation_context", "item_type", "item_id", "item_content", "score", "reason", "is_accepted"],
        {"user": "user_id", "contract": "contract_id"},
        ["user", "contract", "recommendation_type", "recommendation_context", "is_accepted"],
        ["item_content", "reason"],
        "-score",
        soft_delete=False,
        serializer=serialize_recommendation,
    ),
}


@app.get("/")
def root():
    return {"name": "Contract Review FastAPI", "docs": "/docs"}


@app.post("/api/auth/login/")
async def login(request: Request):
    payload = await request.json()
    username = payload.get("username", "")
    password = payload.get("password", "")
    user = db_one(
        "SELECT * FROM users_user WHERE username = %s AND COALESCE(is_deleted, 0) = 0",
        (username,),
    )
    if not user or not verify_django_password(password, user.get("password") or ""):
        raise HTTPException(status_code=401, detail="No active account found with the given credentials")
    if not user.get("is_active"):
        raise HTTPException(status_code=401, detail="User is inactive")
    return {"access": make_token(user, "access"), "refresh": make_token(user, "refresh")}


@app.post("/api/auth/refresh/")
async def refresh_token(request: Request):
    payload = await request.json()
    token = payload.get("refresh")
    if not token:
        raise HTTPException(status_code=400, detail="refresh is required")
    data = decode_token(token)
    if data.get("type") != "refresh":
        raise HTTPException(status_code=401, detail="Invalid refresh token")
    user = get_user_by_id(data.get("user_id"))
    if not user:
        raise HTTPException(status_code=401, detail="User not found")
    return {"access": make_token(user, "access")}


@app.get("/api/users/users/me/")
def me(user: dict[str, Any] = Depends(current_user)):
    return user


@app.get("/api/users/users/")
def list_users(request: Request, user: dict[str, Any] = Depends(current_user)):
    spec = ResourceSpec(
        "users_user",
        USER_FIELDS,
        [],
        [],
        {"department": "department_id"},
        ["role", "department", "is_active"],
        ["username", "email", "real_name"],
        "-created_at",
        serializer=lambda row: serialize_user(get_user_by_id(row["id"]) or row),
    )
    return list_resource(request, spec)


@app.post("/api/users/users/")
async def create_user(request: Request, user: dict[str, Any] = Depends(current_user)):
    payload = await request.json()
    password = payload.pop("password", "")
    data = filter_payload(
        payload,
        ["username", "email", "real_name", "phone", "avatar", "department_id", "role", "reviewer_level", "is_active"],
        {"department": "department_id"},
    )
    data["password"] = make_django_password(password or "123456")
    data["is_staff"] = 0
    data["is_superuser"] = 0
    data["is_deleted"] = 0
    data["created_at"] = now()
    data["updated_at"] = now()
    row = insert_row("users_user", data)
    return serialize_user(get_user_by_id(row["id"]) or row)


@app.get("/api/users/users/{row_id}/")
def retrieve_user(row_id: int, user: dict[str, Any] = Depends(current_user)):
    row = get_user_by_id(row_id)
    if not row or row.get("is_deleted"):
        raise HTTPException(status_code=404, detail="Not found")
    return serialize_user(row)


@app.patch("/api/users/users/{row_id}/")
async def update_user(row_id: int, request: Request, user: dict[str, Any] = Depends(current_user)):
    payload = await request.json()
    password = payload.pop("password", None)
    data = filter_payload(
        payload,
        ["username", "email", "real_name", "phone", "avatar", "department_id", "role", "reviewer_level", "is_active"],
        {"department": "department_id"},
    )
    if password:
        data["password"] = make_django_password(password)
    data["updated_at"] = now()
    update_row("users_user", row_id, data)
    role_ids = payload.get("role_ids")
    if isinstance(role_ids, list):
        db_execute("DELETE FROM users_user_role WHERE user_id = %s", (row_id,))
        for role_id in role_ids:
            insert_row("users_user_role", {"user_id": row_id, "role_id": role_id, "created_at": now()})
    return serialize_user(get_user_by_id(row_id) or {})


@app.delete("/api/users/users/{row_id}/")
def delete_user(row_id: int, user: dict[str, Any] = Depends(current_user)):
    return soft_delete_or_delete("users_user", row_id, True)


@app.post("/api/users/users/{row_id}/assign_roles/")
async def assign_user_roles(row_id: int, request: Request, user: dict[str, Any] = Depends(current_user)):
    payload = await request.json()
    db_execute("DELETE FROM users_user_role WHERE user_id = %s", (row_id,))
    for role_id in payload.get("role_ids", []):
        insert_row("users_user_role", {"user_id": row_id, "role_id": role_id, "created_at": now()})
    return serialize_user(get_user_by_id(row_id) or {})


def mount_generic(prefix: str, spec_key: str):
    spec = SPECS[spec_key]

    async def list_endpoint(request: Request, user: dict[str, Any] = Depends(current_user)):
        return list_resource(request, spec)

    async def create_endpoint(request: Request, user: dict[str, Any] = Depends(current_user)):
        payload = await request.json()
        extra: dict[str, Any] = {}
        if "created_by_id" in spec.fields:
            extra["created_by_id"] = user["id"]
        if spec.table == "recommendations_recommendation" and "user_id" in spec.fields:
            extra["user_id"] = user["id"]
        return create_resource(spec, payload, extra)

    async def retrieve_endpoint(row_id: int, user: dict[str, Any] = Depends(current_user)):
        return get_resource(spec, row_id)

    async def patch_endpoint(row_id: int, request: Request, user: dict[str, Any] = Depends(current_user)):
        payload = await request.json()
        extra_payload = payload.copy()
        if "updated_by_id" in spec.fields:
            extra_payload["updated_by_id"] = user["id"]
        return update_resource(spec, row_id, extra_payload)

    async def put_endpoint(row_id: int, request: Request, user: dict[str, Any] = Depends(current_user)):
        return await patch_endpoint(row_id, request, user)

    async def delete_endpoint(row_id: int, user: dict[str, Any] = Depends(current_user)):
        return soft_delete_or_delete(spec.table, row_id, spec.soft_delete)

    app.add_api_route(prefix, list_endpoint, methods=["GET"])
    app.add_api_route(prefix, create_endpoint, methods=["POST"])
    app.add_api_route(prefix + "{row_id}/", retrieve_endpoint, methods=["GET"])
    app.add_api_route(prefix + "{row_id}/", patch_endpoint, methods=["PATCH"])
    app.add_api_route(prefix + "{row_id}/", put_endpoint, methods=["PUT"])
    app.add_api_route(prefix + "{row_id}/", delete_endpoint, methods=["DELETE"])


@app.get("/api/reviews/ai-model-configs/get_available_models/")
def available_models(provider: str = "siliconflow", user: dict[str, Any] = Depends(current_user)):
    models = [
        {"value": "deepseek-ai/DeepSeek-V3", "label": "DeepSeek-V3"},
        {"value": "Qwen/Qwen2.5-72B-Instruct", "label": "Qwen2.5-72B-Instruct"},
        {"value": "qwen-plus", "label": "Qwen Plus"},
    ]
    return {"provider": provider, "models": models}


@app.get("/api/reviews/focus-configs/by_level/")
def focus_config_by_level(level: str, user: dict[str, Any] = Depends(current_user)):
    row = db_one("SELECT * FROM reviews_review_focus_config WHERE level = %s AND is_active = 1", (level,))
    if not row:
        raise HTTPException(status_code=404, detail="Focus config not found")
    return row


CHAT_SYSTEM_PROMPT = (
    "你是一位专业的合同审核专家助手，擅长解答合同审核、法律合规、风险识别等相关问题。"
    "请用专业、友好、易懂的中文回答用户的问题。"
)

DOCUMENT_OVERVIEW_SYSTEM_PROMPT = (
    "你是文档总结概览助手。请根据用户提供的文档全文生成文档概览。"
    "要求：只输出中文概览正文，不要标题、编号、Markdown或解释说明；"
    "概览必须聚焦文档主题、核心内容、关键事项和结论，控制在100到400字。"
)


def get_default_ai_config() -> dict[str, Any] | None:
    return db_one(
        """
        SELECT *
        FROM reviews_ai_model_config
        WHERE is_default = 1 AND is_active = 1
        ORDER BY id DESC
        LIMIT 1
        """
    ) or db_one(
        """
        SELECT *
        FROM reviews_ai_model_config
        WHERE is_active = 1
        ORDER BY is_default DESC, id DESC
        LIMIT 1
        """
    )


def build_chat_messages(message: str, history: Any) -> list[dict[str, str]]:
    messages = [{"role": "system", "content": CHAT_SYSTEM_PROMPT}]
    if isinstance(history, list):
        for item in history[-10:]:
            if not isinstance(item, dict):
                continue
            role = item.get("role")
            content = str(item.get("content") or "").strip()
            if role in {"user", "assistant"} and content:
                messages.append({"role": role, "content": content})
    messages.append({"role": "user", "content": message})
    return messages


def extract_stream_delta(payload: dict[str, Any]) -> str:
    choices = payload.get("choices") or []
    if choices:
        choice = choices[0] or {}
        delta = choice.get("delta") or {}
        if isinstance(delta, dict):
            content = delta.get("content")
            if content:
                return str(content)
        message = choice.get("message") or {}
        if isinstance(message, dict) and message.get("content"):
            return str(message["content"])
        if choice.get("text"):
            return str(choice["text"])
    for key in ("content", "response", "message", "text", "output_text"):
        value = payload.get(key)
        if isinstance(value, str) and value:
            return value
    return ""


def resolve_ai_model(config: dict[str, Any]) -> str:
    model = str(config.get("default_model") or "").strip()
    available_models = config.get("available_models")
    if not model and isinstance(available_models, list) and available_models:
        model = str(available_models[0] or "").strip()
    return model


def parse_ai_error(response: requests.Response) -> str:
    try:
        error_data = response.json()
        error_message = error_data.get("message") or error_data.get("error") or response.text
        if isinstance(error_message, dict):
            error_message = error_message.get("message") or json.dumps(error_message, ensure_ascii=False)
        return str(error_message)
    except ValueError:
        return response.text[:500]


def call_ai_chat_completion(
    messages: list[dict[str, str]],
    *,
    max_tokens: int | None = None,
    temperature: float | None = None,
) -> tuple[str, str]:
    config = get_default_ai_config()
    if not config or not config.get("api_key"):
        raise HTTPException(status_code=503, detail="AI服务未启用或未配置API密钥")

    model = resolve_ai_model(config)
    if not model:
        raise HTTPException(status_code=503, detail="AI服务未配置默认模型")

    api_base_url = str(config.get("api_base_url") or "").rstrip("/")
    if not api_base_url:
        raise HTTPException(status_code=503, detail="AI服务未配置API基础地址")

    timeout = int(config.get("timeout") or 30)
    request_data = {
        "model": model,
        "messages": messages,
        "temperature": temperature if temperature is not None else float(config.get("temperature") or 0.7),
        "max_tokens": max_tokens if max_tokens is not None else int(config.get("max_tokens") or 2000),
        "stream": False,
    }
    headers = {
        "Authorization": f"Bearer {config['api_key']}",
        "Content-Type": "application/json",
    }

    try:
        response = requests.post(
            f"{api_base_url}/chat/completions",
            headers=headers,
            json=request_data,
            timeout=(10, timeout),
        )
    except requests.exceptions.Timeout as exc:
        raise HTTPException(status_code=504, detail="AI接口调用超时，请稍后重试") from exc
    except requests.exceptions.ConnectionError as exc:
        raise HTTPException(status_code=502, detail="无法连接到AI服务，请检查API地址和网络连接") from exc

    if response.status_code != 200:
        raise HTTPException(status_code=502, detail=f"AI接口调用失败：{parse_ai_error(response)}")

    try:
        payload = response.json()
    except ValueError as exc:
        raise HTTPException(status_code=502, detail="AI接口返回格式无法解析") from exc

    content = extract_stream_delta(payload).strip()
    if not content:
        raise HTTPException(status_code=502, detail="AI接口未返回有效内容")
    return content, model


def normalize_document_overview(text: str) -> str:
    overview = re.sub(r"\s+", " ", text or "").strip()
    if len(overview) <= 400:
        return overview
    clipped = overview[:400]
    for separator in ("。", "！", "？", "；"):
        index = clipped.rfind(separator)
        if index >= 100:
            return clipped[: index + 1].strip()
    return clipped.strip()


@app.post("/api/portal/ai/document-overview/")
@app.post("/api/portal/ai/document-overview-chat/")
async def portal_document_overview_chat(request: Request):
    payload = await request.json()
    content = str(
        payload.get("content")
        or payload.get("document_text")
        or payload.get("message")
        or ""
    ).strip()
    if not content:
        raise HTTPException(status_code=400, detail="文档内容不能为空")

    document_name = str(payload.get("document_name") or payload.get("filename") or "").strip()
    user_content = "请为以下文档生成100到400字的概览。"
    if document_name:
        user_content += f"\n文档名称：{document_name}"
    user_content += f"\n\n文档全文：\n{content}"

    response_text, model = call_ai_chat_completion(
        [
            {"role": "system", "content": DOCUMENT_OVERVIEW_SYSTEM_PROMPT},
            {"role": "user", "content": user_content},
        ],
        max_tokens=700,
        temperature=0.2,
    )
    overview = normalize_document_overview(response_text)
    return {"success": True, "overview": overview, "response": overview, "model": model}


def stream_chat_chunks(message: str, history: Any):
    config = get_default_ai_config()
    if not config or not config.get("api_key"):
        yield "抱歉，AI服务未启用或未配置。请管理员在AI模型配置中设置API密钥和默认模型。"
        return

    model = resolve_ai_model(config)
    if not model:
        yield "抱歉，AI服务未配置默认模型。请管理员在AI模型配置中设置默认模型。"
        return

    api_base_url = str(config.get("api_base_url") or "").rstrip("/")
    if not api_base_url:
        yield "抱歉，AI服务未配置API基础地址。"
        return

    timeout = int(config.get("timeout") or 30)
    request_data = {
        "model": model,
        "messages": build_chat_messages(message, history),
        "temperature": float(config.get("temperature") or 0.7),
        "max_tokens": int(config.get("max_tokens") or 2000),
        "stream": True,
    }
    headers = {
        "Authorization": f"Bearer {config['api_key']}",
        "Content-Type": "application/json",
    }

    try:
        with requests.post(
            f"{api_base_url}/chat/completions",
            headers=headers,
            json=request_data,
            stream=True,
            timeout=(10, timeout),
        ) as response:
            if response.status_code != 200:
                try:
                    error_data = response.json()
                    error_message = error_data.get("message") or error_data.get("error") or response.text
                    if isinstance(error_message, dict):
                        error_message = error_message.get("message") or json.dumps(error_message, ensure_ascii=False)
                except ValueError:
                    error_message = response.text[:500]
                yield f"抱歉，AI接口调用失败：{error_message}"
                return

            content_type = response.headers.get("content-type", "")
            if "text/event-stream" not in content_type:
                try:
                    payload = response.json()
                    content = extract_stream_delta(payload)
                    yield content or "抱歉，AI接口返回格式无法解析。"
                except ValueError:
                    yield response.text
                return

            for raw_line in response.iter_lines(chunk_size=1, decode_unicode=True):
                if not raw_line:
                    continue
                line = raw_line.decode("utf-8", errors="ignore") if isinstance(raw_line, bytes) else raw_line
                line = line.strip()
                if not line.startswith("data:"):
                    continue
                data = line[5:].strip()
                if not data or data == "[DONE]":
                    if data == "[DONE]":
                        break
                    continue
                try:
                    payload = json.loads(data)
                except json.JSONDecodeError:
                    continue
                delta = extract_stream_delta(payload)
                if delta:
                    yield delta
    except requests.exceptions.Timeout:
        yield "抱歉，AI接口调用超时，请稍后重试。"
    except requests.exceptions.ConnectionError:
        yield "抱歉，无法连接到AI服务，请检查API地址和网络连接。"
    except Exception as exc:
        yield f"抱歉，AI对话失败：{exc}"


@app.post("/api/reviews/ai-model-configs/chat/")
async def chat_endpoint(request: Request, user: dict[str, Any] = Depends(current_user)):
    payload = await request.json()
    message = str(payload.get("message") or "").strip()
    if not message:
        raise HTTPException(status_code=400, detail="消息内容不能为空")
    response_text = "".join(stream_chat_chunks(message, payload.get("history", [])))
    return {"response": response_text, "model": (get_default_ai_config() or {}).get("default_model")}


@app.post("/api/reviews/ai-model-configs/stream-chat")
@app.post("/api/reviews/ai-model-configs/stream-chat/")
async def stream_chat_endpoint(request: Request, user: dict[str, Any] = Depends(current_user)):
    payload = await request.json()
    message = str(payload.get("message") or "").strip()
    if not message:
        raise HTTPException(status_code=400, detail="消息内容不能为空")
    return StreamingResponse(
        stream_chat_chunks(message, payload.get("history", [])),
        media_type="text/plain; charset=utf-8",
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no",
        },
    )


mount_generic("/api/users/departments/", "departments")
mount_generic("/api/users/roles/", "roles")
mount_generic("/api/users/permissions/", "permissions")
mount_generic("/api/contracts/templates/", "templates")
mount_generic("/api/reviews/focus-configs/", "review_focus_configs")
mount_generic("/api/reviews/ai-model-configs/", "ai_model_configs")
mount_generic("/api/rules/rules/", "rules")
mount_generic("/api/rules/matches/", "matches")
mount_generic("/api/knowledge/entities/", "entities")
mount_generic("/api/knowledge/relations/", "relations")
mount_generic("/api/knowledge/regulations/", "regulations")
mount_generic("/api/knowledge/cases/", "cases")
mount_generic("/api/clauses/clauses/", "clauses")
mount_generic("/api/risks/risks/", "risks")
mount_generic("/api/comparisons/tasks/", "comparison_tasks")
mount_generic("/api/comparisons/diffs/", "comparison_diffs")
mount_generic("/api/recommendations/recommendations/", "recommendations")


@app.post("/api/clauses/clauses/{row_id}/confirm/")
def confirm_clause(row_id: int, user: dict[str, Any] = Depends(current_user)):
    update_row(
        "clauses_contract_clause",
        row_id,
        {
            "is_confirmed": 1,
            "confirmed_by_id": user["id"],
            "confirmed_at": now(),
            "updated_at": now(),
        },
    )
    return get_resource(SPECS["clauses"], row_id)


@app.post("/api/risks/risks/{row_id}/handle/")
def handle_risk(row_id: int, user: dict[str, Any] = Depends(current_user)):
    update_row(
        "risks_risk_identification",
        row_id,
        {
            "status": "handled",
            "handled_by_id": user["id"],
            "handled_at": now(),
        },
    )
    return get_resource(SPECS["risks"], row_id)


@app.post("/api/recommendations/recommendations/{row_id}/accept/")
def accept_recommendation(row_id: int, user: dict[str, Any] = Depends(current_user)):
    update_row("recommendations_recommendation", row_id, {"is_accepted": 1})
    return get_resource(SPECS["recommendations"], row_id)


@app.post("/api/recommendations/recommendations/{row_id}/reject/")
def reject_recommendation(row_id: int, user: dict[str, Any] = Depends(current_user)):
    update_row("recommendations_recommendation", row_id, {"is_accepted": 0})
    return get_resource(SPECS["recommendations"], row_id)


PORTAL_DOCUMENT_ROOT = MEDIA_ROOT / "portal" / "documents"
PORTAL_ALLOWED_EXTENSIONS = {".docx", ".pdf"}
PORTAL_MAX_UPLOAD_SIZE = int(os.getenv("PORTAL_MAX_UPLOAD_SIZE", str(20 * 1024 * 1024)))
PORTAL_DOCX_CHARS_PER_PAGE = int(os.getenv("PORTAL_DOCX_CHARS_PER_PAGE", "1800"))
PORTAL_DOCUMENT_ID_RE = re.compile(r"^[a-f0-9]{32}$")


def sanitize_filename(filename: str) -> str:
    name = Path(filename or "").name.strip()
    if not name:
        return "未命名文档"
    return re.sub(r"[\\/:*?\"<>|]+", "_", name)[:180]


def count_preview_words(text: str) -> int:
    return len(re.findall(r"[\u4e00-\u9fff]|[A-Za-z0-9]+(?:[-_][A-Za-z0-9]+)*", text or ""))


def validate_portal_document_id(document_id: str) -> None:
    if not PORTAL_DOCUMENT_ID_RE.match(document_id):
        raise HTTPException(status_code=404, detail="文档不存在")


def portal_document_dir(document_id: str) -> Path:
    validate_portal_document_id(document_id)
    return PORTAL_DOCUMENT_ROOT / document_id


def portal_metadata_path(document_id: str) -> Path:
    return portal_document_dir(document_id) / "metadata.json"


def load_portal_metadata(document_id: str) -> dict[str, Any]:
    metadata_path = portal_metadata_path(document_id)
    if not metadata_path.exists():
        raise HTTPException(status_code=404, detail="文档不存在")
    try:
        return json.loads(metadata_path.read_text(encoding="utf-8"))
    except json.JSONDecodeError as exc:
        raise HTTPException(status_code=500, detail="文档元数据损坏") from exc


def save_portal_metadata(metadata: dict[str, Any]) -> None:
    metadata_path = portal_metadata_path(metadata["id"])
    metadata_path.write_text(json.dumps(metadata, ensure_ascii=False, indent=2), encoding="utf-8")


def validate_uploaded_file_signature(file_path: Path, extension: str) -> None:
    if extension == ".pdf":
        with file_path.open("rb") as source:
            if source.read(4) != b"%PDF":
                raise HTTPException(status_code=400, detail="文件内容不是有效的 PDF")
        return
    if extension == ".docx":
        if not zipfile.is_zipfile(file_path):
            raise HTTPException(status_code=400, detail="文件内容不是有效的 DOCX")
        with zipfile.ZipFile(file_path) as archive:
            names = set(archive.namelist())
            if "[Content_Types].xml" not in names or "word/document.xml" not in names:
                raise HTTPException(status_code=400, detail="文件内容不是有效的 DOCX")


def docx_table_to_block(table: Any) -> dict[str, str]:
    rows: list[str] = []
    html_rows: list[str] = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        if any(cells):
            rows.append(" | ".join(cells))
        html_cells = "".join(f"<td>{html.escape(cell)}</td>" for cell in cells)
        html_rows.append(f"<tr>{html_cells}</tr>")
    text = "\n".join(rows)
    return {
        "text": text,
        "html": f"<table class=\"docx-preview-table\"><tbody>{''.join(html_rows)}</tbody></table>",
    }


def extract_docx_preview(file_path: Path) -> dict[str, Any]:
    if docx is None:
        raise HTTPException(status_code=500, detail="服务端缺少 python-docx 依赖")
    document = docx.Document(file_path)
    blocks: list[dict[str, str]] = []
    text_parts: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        paragraph_xml = paragraph._element.xml
        has_page_break = 'w:type="page"' in paragraph_xml or "lastRenderedPageBreak" in paragraph_xml
        if not text:
            if has_page_break:
                blocks.append({"text": "", "html": "", "page_break": "1"})
            continue
        text_parts.append(text)
        style_name = (getattr(paragraph.style, "name", "") or "").lower()
        tag = "h2" if "heading" in style_name or "标题" in style_name else "p"
        blocks.append({"text": text, "html": f"<{tag}>{html.escape(text)}</{tag}>"})
        if has_page_break:
            blocks.append({"text": "", "html": "", "page_break": "1"})

    for table in document.tables:
        block = docx_table_to_block(table)
        if block["text"]:
            text_parts.append(block["text"])
            blocks.append(block)

    pages: list[dict[str, Any]] = []
    current_html: list[str] = []
    current_text: list[str] = []
    current_chars = 0
    for block in blocks:
        if block.get("page_break"):
            if current_html:
                pages.append(
                    {
                        "page_number": len(pages) + 1,
                        "type": "docx_html",
                        "html": "".join(current_html),
                        "text": "\n".join(current_text),
                    }
                )
                current_html = []
                current_text = []
                current_chars = 0
            continue
        block_len = len(block["text"])
        if current_html and current_chars + block_len > PORTAL_DOCX_CHARS_PER_PAGE:
            pages.append(
                {
                    "page_number": len(pages) + 1,
                    "type": "docx_html",
                    "html": "".join(current_html),
                    "text": "\n".join(current_text),
                }
            )
            current_html = []
            current_text = []
            current_chars = 0
        current_html.append(block["html"])
        current_text.append(block["text"])
        current_chars += block_len

    if current_html:
        pages.append(
            {
                "page_number": len(pages) + 1,
                "type": "docx_html",
                "html": "".join(current_html),
                "text": "\n".join(current_text),
            }
        )
    if not pages:
        pages.append({"page_number": 1, "type": "docx_html", "html": "<p>该文档没有可预览文本。</p>", "text": ""})

    full_text = "\n".join(text_parts)
    title = ""
    try:
        title = (document.core_properties.title or "").strip()
    except Exception:
        title = ""
    if not title:
        title = next((part for part in text_parts if part), "")

    return {
        "title": title[:120],
        "page_count": len(pages),
        "word_count": count_preview_words(full_text),
        "pages": pages,
    }


def extract_pdf_preview(document_id: str, file_path: Path) -> dict[str, Any]:
    if fitz is None:
        raise HTTPException(status_code=500, detail="服务端缺少 PyMuPDF 依赖")
    text_parts: list[str] = []
    pages: list[dict[str, Any]] = []
    with fitz.open(file_path) as pdf:
        metadata = pdf.metadata or {}
        for page_index in range(len(pdf)):
            page = pdf[page_index]
            text = page.get_text("text").strip()
            text_parts.append(text)
            pages.append(
                {
                    "page_number": page_index + 1,
                    "type": "pdf_image",
                    "image_url": f"/api/portal/documents/{document_id}/pages/{page_index + 1}.png",
                    "text": text,
                }
            )
        title = (metadata.get("title") or "").strip()
        if not title:
            first_lines = [line.strip() for line in (text_parts[0] if text_parts else "").splitlines() if line.strip()]
            title = first_lines[0] if first_lines else ""

    if not pages:
        pages.append({"page_number": 1, "type": "pdf_image", "image_url": f"/api/portal/documents/{document_id}/pages/1.png", "text": ""})

    full_text = "\n".join(text_parts)
    return {
        "title": title[:120],
        "page_count": len(pages),
        "word_count": count_preview_words(full_text),
        "pages": pages,
    }


def build_portal_preview(document_id: str, file_path: Path, extension: str) -> dict[str, Any]:
    if extension == ".pdf":
        return extract_pdf_preview(document_id, file_path)
    if extension == ".docx":
        return extract_docx_preview(file_path)
    raise HTTPException(status_code=400, detail="仅支持 DOCX、PDF 文件")


def portal_url(request: Request, path: str) -> str:
    base_url = str(request.base_url).rstrip("/")
    return f"{base_url}{path}"


def serialize_portal_document(request: Request, metadata: dict[str, Any]) -> dict[str, Any]:
    document_id = metadata["id"]
    file_url = f"/api/portal/documents/{document_id}/file/"
    preview_url = f"/api/portal/documents/{document_id}/preview/"
    pages: list[dict[str, Any]] = []
    for page in metadata.get("pages", []):
        page_data = page.copy()
        if page_data.get("image_url"):
            page_data["absolute_image_url"] = portal_url(request, page_data["image_url"])
        pages.append(page_data)
    return {
        "id": document_id,
        "name": metadata["original_name"],
        "title": metadata.get("title") or metadata["original_name"],
        "extension": metadata["extension"],
        "content_type": metadata["content_type"],
        "size": metadata["size"],
        "page_count": metadata["page_count"],
        "word_count": metadata["word_count"],
        "created_at": metadata["created_at"],
        "file_url": file_url,
        "preview_url": preview_url,
        "absolute_file_url": portal_url(request, file_url),
        "absolute_preview_url": portal_url(request, preview_url),
        "pages": pages,
    }


@app.post("/api/portal/documents/upload/")
async def upload_portal_document(request: Request, file: UploadFile = File(...)):
    original_name = sanitize_filename(file.filename or "")
    extension = Path(original_name).suffix.lower()
    if extension not in PORTAL_ALLOWED_EXTENSIONS:
        raise HTTPException(status_code=400, detail="仅支持上传 DOCX、PDF 文件")

    document_id = uuid.uuid4().hex
    document_dir = portal_document_dir(document_id)
    document_dir.mkdir(parents=True, exist_ok=True)
    source_path = document_dir / f"source{extension}"

    size = 0
    try:
        with source_path.open("wb") as destination:
            while True:
                chunk = await file.read(1024 * 1024)
                if not chunk:
                    break
                size += len(chunk)
                if size > PORTAL_MAX_UPLOAD_SIZE:
                    raise HTTPException(status_code=400, detail="文件大小不能超过 20MB")
                destination.write(chunk)
        if size == 0:
            raise HTTPException(status_code=400, detail="不能上传空文件")
        validate_uploaded_file_signature(source_path, extension)
        preview = build_portal_preview(document_id, source_path, extension)
        content_type = (
            "application/pdf"
            if extension == ".pdf"
            else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        metadata = {
            "id": document_id,
            "original_name": original_name,
            "extension": extension[1:],
            "content_type": content_type,
            "size": size,
            "source_file": source_path.name,
            "title": preview.get("title") or original_name,
            "page_count": preview["page_count"],
            "word_count": preview["word_count"],
            "pages": preview["pages"],
            "created_at": datetime.now().isoformat(timespec="seconds"),
        }
        save_portal_metadata(metadata)
        return {"success": True, "message": "文件上传成功", "document": serialize_portal_document(request, metadata)}
    except HTTPException:
        shutil.rmtree(document_dir, ignore_errors=True)
        raise
    except Exception as exc:
        shutil.rmtree(document_dir, ignore_errors=True)
        raise HTTPException(status_code=500, detail=f"文件处理失败: {exc}") from exc
    finally:
        await file.close()


@app.get("/api/portal/documents/{document_id}/preview/")
def get_portal_document_preview(document_id: str, request: Request):
    metadata = load_portal_metadata(document_id)
    return {"success": True, "document": serialize_portal_document(request, metadata)}


@app.get("/api/portal/documents/{document_id}/file/")
def get_portal_document_file(document_id: str):
    metadata = load_portal_metadata(document_id)
    file_path = portal_document_dir(document_id) / metadata["source_file"]
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="源文件不存在")
    disposition = f"inline; filename*=UTF-8''{quote(metadata['original_name'])}"
    return FileResponse(
        file_path,
        media_type=metadata["content_type"],
        headers={"Content-Disposition": disposition},
    )


@app.get("/api/portal/documents/{document_id}/pages/{page_number}.png")
def get_portal_pdf_page_image(document_id: str, page_number: int):
    metadata = load_portal_metadata(document_id)
    if metadata["extension"] != "pdf":
        raise HTTPException(status_code=404, detail="该文档没有 PDF 页图")
    page_count = int(metadata.get("page_count") or 0)
    if page_number < 1 or page_number > page_count:
        raise HTTPException(status_code=404, detail="页码不存在")
    if fitz is None:
        raise HTTPException(status_code=500, detail="服务端缺少 PyMuPDF 依赖")

    document_dir = portal_document_dir(document_id)
    pages_dir = document_dir / "pages"
    pages_dir.mkdir(exist_ok=True)
    image_path = pages_dir / f"{page_number}.png"
    if not image_path.exists():
        source_path = document_dir / metadata["source_file"]
        if not source_path.exists():
            raise HTTPException(status_code=404, detail="源文件不存在")
        with fitz.open(source_path) as pdf:
            page = pdf[page_number - 1]
            pixmap = page.get_pixmap(matrix=fitz.Matrix(1.8, 1.8), alpha=False)
            pixmap.save(image_path)
    return FileResponse(image_path, media_type="image/png")


@app.post("/api/recommendations/recommendations/recommend_clauses/")
async def recommend_clauses(request: Request, user: dict[str, Any] = Depends(current_user)):
    payload = await request.json()
    contract_id = payload.get("contract_id")
    if not contract_id:
        raise HTTPException(status_code=400, detail="contract_id不能为空")
    contract = db_one("SELECT id FROM contracts_contract WHERE id = %s AND COALESCE(is_deleted, 0) = 0", (contract_id,))
    if not contract:
        raise HTTPException(status_code=404, detail="合同不存在")
    clauses = db_query(
        """
        SELECT id, clause_title, clause_content, confidence
        FROM clauses_contract_clause
        WHERE contract_id = %s
        ORDER BY confidence DESC, created_at DESC
        LIMIT 20
        """,
        (contract_id,),
    )
    recommendations = [
        {
            "item_type": "clause",
            "item_id": row["id"],
            "item_content": row.get("clause_content") or row.get("clause_title") or "",
            "score": row.get("confidence") or 0,
            "reason": "基于当前合同已识别条款推荐",
        }
        for row in clauses
    ]
    return {"success": True, "recommendations": recommendations, "count": len(recommendations)}


@app.post("/api/recommendations/recommendations/recommend_templates/")
async def recommend_templates(request: Request, user: dict[str, Any] = Depends(current_user)):
    payload = await request.json()
    contract_type = payload.get("contract_type")
    industry = payload.get("industry", "")
    if not contract_type:
        raise HTTPException(status_code=400, detail="contract_type不能为空")
    rows = db_query(
        """
        SELECT id, name, content, usage_count
        FROM contracts_template
        WHERE contract_type = %s
          AND COALESCE(is_deleted, 0) = 0
          AND (%s = '' OR industry = %s OR industry = '')
        ORDER BY usage_count DESC, created_at DESC
        LIMIT 20
        """,
        (contract_type, industry, industry),
    )
    recommendations = [
        {
            "item_type": "template",
            "item_id": row["id"],
            "item_content": row.get("name") or row.get("content") or "",
            "score": row.get("usage_count") or 0,
            "reason": "基于合同类型和行业匹配模板",
        }
        for row in rows
    ]
    return {"success": True, "recommendations": recommendations, "count": len(recommendations)}


@app.post("/api/recommendations/recommendations/recommend_risk_responses/")
async def recommend_risk_responses(request: Request, user: dict[str, Any] = Depends(current_user)):
    payload = await request.json()
    contract_id = payload.get("contract_id")
    if not contract_id:
        raise HTTPException(status_code=400, detail="contract_id不能为空")
    contract = db_one("SELECT id FROM contracts_contract WHERE id = %s AND COALESCE(is_deleted, 0) = 0", (contract_id,))
    if not contract:
        raise HTTPException(status_code=404, detail="合同不存在")
    rows = db_query(
        """
        SELECT id, risk_description, suggestion, risk_level
        FROM risks_risk_identification
        WHERE contract_id = %s
        ORDER BY risk_level DESC, created_at DESC
        LIMIT 20
        """,
        (contract_id,),
    )
    recommendations = [
        {
            "item_type": "risk_response",
            "item_id": row["id"],
            "item_content": row.get("suggestion") or row.get("risk_description") or "",
            "score": {"high": 100, "medium": 70, "low": 40}.get(str(row.get("risk_level")), 0),
            "reason": "基于当前合同风险识别结果推荐",
        }
        for row in rows
    ]
    return {"success": True, "recommendations": recommendations, "count": len(recommendations)}


@app.get("/api/users/audit-logs/")
def list_audit_logs(request: Request, user: dict[str, Any] = Depends(current_user)):
    spec = ResourceSpec(
        "users_audit_log",
        ["id", "user_id", "action", "resource_type", "resource_id", "ip_address", "user_agent", "request_data", "response_data", "status", "error_message", "created_at"],
        [],
        [],
        {"user": "user_id"},
        ["user", "action", "status", "resource_type"],
        ["action", "resource_type", "error_message"],
        "-created_at",
        soft_delete=False,
        serializer=lambda row: {**row, "user": row.get("user_id"), "user_name": lookup_name("users_user", row.get("user_id")) or "-"},
    )
    return list_resource(request, spec)


@app.get("/api/users/audit-logs/{row_id}/")
def get_audit_log(row_id: int, user: dict[str, Any] = Depends(current_user)):
    row = db_one("SELECT * FROM users_audit_log WHERE id = %s", (row_id,))
    if not row:
        raise HTTPException(status_code=404, detail="Not found")
    row["user"] = row.get("user_id")
    row["user_name"] = lookup_name("users_user", row.get("user_id")) or "-"
    return row


@app.post("/api/users/roles/{row_id}/assign_permissions/")
async def assign_permissions(row_id: int, request: Request, user: dict[str, Any] = Depends(current_user)):
    payload = await request.json()
    db_execute("DELETE FROM users_role_permission WHERE role_id = %s", (row_id,))
    for permission_id in payload.get("permission_ids", []):
        insert_row("users_role_permission", {"role_id": row_id, "permission_id": permission_id, "created_at": now()})
    return serialize_role(db_one("SELECT * FROM users_role WHERE id = %s", (row_id,)) or {})


@app.get("/api/contracts/contracts/")
def list_contracts(request: Request, user: dict[str, Any] = Depends(current_user)):
    return list_resource(request, SPECS["contracts"])


@app.post("/api/contracts/contracts/")
async def create_contract(request: Request, user: dict[str, Any] = Depends(current_user)):
    payload = await request.json()
    contract_no = f"CT{datetime.now().strftime('%Y%m%d')}{uuid.uuid4().hex[:8].upper()}"
    extra = {
        "contract_no": contract_no,
        "drafter_id": user["id"],
        "status": payload.get("status", "draft"),
        "current_version": 1,
        "is_deleted": 0,
    }
    return create_resource(SPECS["contracts"], payload, extra)


@app.get("/api/contracts/contracts/{row_id}/")
def retrieve_contract(row_id: int, user: dict[str, Any] = Depends(current_user)):
    return get_resource(SPECS["contracts"], row_id)


@app.patch("/api/contracts/contracts/{row_id}/")
async def update_contract(row_id: int, request: Request, user: dict[str, Any] = Depends(current_user)):
    payload = await request.json()
    return update_resource(SPECS["contracts"], row_id, payload)


@app.delete("/api/contracts/contracts/{row_id}/")
def delete_contract(row_id: int, user: dict[str, Any] = Depends(current_user)):
    return soft_delete_or_delete("contracts_contract", row_id, True)


@app.get("/api/contracts/contracts/{row_id}/versions/")
def contract_versions(row_id: int, user: dict[str, Any] = Depends(current_user)):
    return db_query(
        "SELECT * FROM contracts_contract_version WHERE contract_id = %s AND COALESCE(is_deleted, 0) = 0 ORDER BY version DESC",
        (row_id,),
    )


@app.post("/api/contracts/contracts/{row_id}/create_version/")
async def create_contract_version(row_id: int, request: Request, user: dict[str, Any] = Depends(current_user)):
    payload = await request.json()
    contract = db_one("SELECT * FROM contracts_contract WHERE id = %s", (row_id,))
    if not contract:
        raise HTTPException(status_code=404, detail="Contract not found")
    new_version = int(contract.get("current_version") or 1) + 1
    version = insert_row(
        "contracts_contract_version",
        {
            "contract_id": row_id,
            "version": new_version,
            "content": encode_db_value("content", payload.get("content", contract.get("content"))),
            "file_path": payload.get("file_path", contract.get("file_path") or ""),
            "change_summary": payload.get("change_summary") or f"Edit contract - {now()}",
            "changed_by_id": user["id"],
            "is_deleted": 0,
            "created_at": now(),
        },
    )
    update_row("contracts_contract", row_id, {"current_version": new_version, "updated_at": now()})
    return version


@app.post("/api/contracts/contracts/{row_id}/rollback/")
async def rollback_contract(row_id: int, request: Request, user: dict[str, Any] = Depends(current_user)):
    payload = await request.json()
    version_number = payload.get("version")
    if not version_number:
        raise HTTPException(status_code=400, detail="version is required")
    version = db_one(
        "SELECT * FROM contracts_contract_version WHERE contract_id = %s AND version = %s AND COALESCE(is_deleted, 0) = 0",
        (row_id, version_number),
    )
    if not version:
        raise HTTPException(status_code=404, detail="Version not found")
    contract = db_one("SELECT * FROM contracts_contract WHERE id = %s", (row_id,))
    new_version = int(contract.get("current_version") or 1) + 1
    update_row(
        "contracts_contract",
        row_id,
        {"content": encode_db_value("content", version.get("content")), "file_path": version.get("file_path") or "", "current_version": new_version, "updated_at": now()},
    )
    insert_row(
        "contracts_contract_version",
        {
            "contract_id": row_id,
            "version": new_version,
            "content": encode_db_value("content", version.get("content")),
            "file_path": version.get("file_path") or "",
            "change_summary": f"Rollback to version {version_number}",
            "changed_by_id": user["id"],
            "is_deleted": 0,
            "created_at": now(),
        },
    )
    return {"message": f"Rolled back to version {version_number}", "new_version": new_version}


@app.post("/api/contracts/templates/{row_id}/use/")
def use_template(row_id: int, user: dict[str, Any] = Depends(current_user)):
    db_execute("UPDATE contracts_template SET usage_count = COALESCE(usage_count, 0) + 1 WHERE id = %s", (row_id,))
    return {"message": "Template usage count updated"}


@app.post("/api/contracts/contracts/generate_content/")
async def generate_contract_content(request: Request, user: dict[str, Any] = Depends(current_user)):
    payload = await request.json()
    contract_type = payload.get("contract_type") or "service"
    info = payload.get("basic_info") or {}
    text = (
        f"{contract_type} contract\n\n"
        f"Party A: {info.get('party_a', '')}\n"
        f"Party B: {info.get('party_b', '')}\n"
        f"Subject: {info.get('subject', '')}\n"
        f"Amount: {info.get('amount', '')}\n"
        "Please supplement detailed terms, payment, delivery, liability, confidentiality, and dispute resolution."
    )
    return {"success": True, "content": {"text": text, "html": text.replace("\n", "<br>")}}


@app.post("/api/contracts/upload/")
async def upload_contract(file: UploadFile = File(...), user: dict[str, Any] = Depends(current_user)):
    suffix = Path(file.filename or "").suffix.lower()
    if suffix not in {".doc", ".docx", ".pdf"}:
        raise HTTPException(status_code=400, detail="Unsupported file type")
    upload_dir = MEDIA_ROOT / "contracts" / "uploads"
    upload_dir.mkdir(parents=True, exist_ok=True)
    filename = f"{uuid.uuid4().hex}{suffix}"
    full_path = upload_dir / filename
    with full_path.open("wb") as target:
        shutil.copyfileobj(file.file, target)
    relative_path = f"contracts/uploads/{filename}"
    return {"file_path": relative_path, "content": parse_uploaded_file(full_path, suffix), "message": "Upload success"}


def parse_uploaded_file(path: Path, suffix: str) -> dict[str, Any]:
    content = {"text": "", "html": "", "title": "", "metadata": {}}
    try:
        if suffix == ".docx":
            import docx

            doc = docx.Document(str(path))
            paragraphs = [para.text for para in doc.paragraphs if para.text]
            content["text"] = "\n".join(paragraphs)
            content["html"] = "<br>".join(paragraphs)
            content["title"] = paragraphs[0][:200] if paragraphs else ""
        elif suffix == ".pdf":
            import fitz

            pdf = fitz.open(str(path))
            pages = [page.get_text() for page in pdf]
            content["text"] = "\n".join(pages)
            content["html"] = content["text"].replace("\n", "<br>")
            content["title"] = content["text"].strip().splitlines()[0][:200] if content["text"].strip() else ""
    except Exception as exc:
        content["metadata"]["parse_error"] = str(exc)
    return content


@app.get("/api/reviews/tasks/")
def list_review_tasks(request: Request, user: dict[str, Any] = Depends(current_user)):
    return list_resource(request, SPECS["review_tasks"])


@app.post("/api/reviews/tasks/")
async def create_review_task(request: Request, user: dict[str, Any] = Depends(current_user)):
    payload = await request.json()
    extra = {"created_by_id": user["id"], "status": payload.get("status", "pending"), "task_type": payload.get("task_type", "auto")}
    return create_resource(SPECS["review_tasks"], payload, extra)


@app.get("/api/reviews/tasks/reviewers/")
def list_reviewers(level: str | None = None, user: dict[str, Any] = Depends(current_user)):
    where = "role = 'reviewer' AND is_active = 1 AND COALESCE(is_deleted, 0) = 0"
    params: tuple[Any, ...] = ()
    if level:
        where += " AND reviewer_level = %s"
        params = (level,)
    rows = db_query(f"SELECT id, username, real_name, email, reviewer_level FROM users_user WHERE {where}", params)
    grouped: dict[str, list[dict[str, Any]]] = {}
    for row in rows:
        key = row.get("reviewer_level") or "unassigned"
        grouped.setdefault(key, []).append({**row, "level": row.get("reviewer_level")})
    return {"reviewers": rows, "grouped": grouped}


@app.get("/api/reviews/tasks/{row_id}/")
def retrieve_review_task(row_id: int, user: dict[str, Any] = Depends(current_user)):
    return get_resource(SPECS["review_tasks"], row_id)


@app.patch("/api/reviews/tasks/{row_id}/")
async def update_review_task(row_id: int, request: Request, user: dict[str, Any] = Depends(current_user)):
    payload = await request.json()
    return update_resource(SPECS["review_tasks"], row_id, payload)


@app.delete("/api/reviews/tasks/{row_id}/")
def delete_review_task(row_id: int, user: dict[str, Any] = Depends(current_user)):
    return soft_delete_or_delete("reviews_review_task", row_id, False)


@app.post("/api/reviews/tasks/{row_id}/start/")
def start_review_task(row_id: int, user: dict[str, Any] = Depends(current_user)):
    task = db_one("SELECT * FROM reviews_review_task WHERE id = %s", (row_id,))
    if not task:
        raise HTTPException(status_code=404, detail="Task not found")
    if task.get("status") not in {"pending", "ai_completed", "manual_reviewing"}:
        raise HTTPException(status_code=400, detail="Task status does not allow start")
    update_row(
        "reviews_review_task",
        row_id,
        {
            "status": "manual_reviewing",
            "started_at": task.get("started_at") or now(),
            "progress": encode_db_value(
                "progress",
                {"current_step": "review ready", "progress": 100, "message": "Ready for manual review", "steps": []},
            ),
            "updated_at": now(),
        },
    )
    ensure_review_result(task)
    return {"message": "Review task started", "task": get_resource(SPECS["review_tasks"], row_id)}


def ensure_review_result(task: dict[str, Any]) -> dict[str, Any]:
    result = db_one("SELECT * FROM reviews_review_result WHERE review_task_id = %s", (task["id"],))
    if result:
        return result
    return insert_row(
        "reviews_review_result",
        {
            "review_task_id": task["id"],
            "contract_id": task["contract_id"],
            "overall_score": 85,
            "risk_level": "low",
            "risk_count": 0,
            "summary": "Review initialized.",
            "report_path": "",
            "report_format": "",
            "review_data": encode_db_value("review_data", {"risk_overview": {"risk_count": 0, "risk_level": "low"}}),
            "created_at": now(),
        },
    )


@app.post("/api/reviews/tasks/{row_id}/complete_manually/")
def complete_review_task(row_id: int, user: dict[str, Any] = Depends(current_user)):
    update_row("reviews_review_task", row_id, {"status": "completed", "completed_at": now(), "updated_at": now()})
    return {"message": "Task completed"}


@app.post("/api/reviews/tasks/{row_id}/submit_review/")
async def submit_review(row_id: int, request: Request, user: dict[str, Any] = Depends(current_user)):
    payload = await request.json()
    task = db_one("SELECT * FROM reviews_review_task WHERE id = %s", (row_id,))
    if not task:
        raise HTTPException(status_code=404, detail="Task not found")
    result = ensure_review_result(task)
    for opinion in payload.get("opinions", []):
        insert_row(
            "reviews_review_opinion",
            {
                "review_result_id": result["id"],
                "reviewer_id": user["id"],
                "clause_id": opinion.get("clause_id", ""),
                "clause_content": opinion.get("clause_content", ""),
                "opinion_type": opinion.get("opinion_type", "suggestion"),
                "risk_level": opinion.get("risk_level", "low"),
                "opinion_content": opinion.get("opinion_content", ""),
                "legal_basis": opinion.get("legal_basis", ""),
                "suggestion": opinion.get("suggestion", ""),
                "status": "pending",
                "is_deleted": 0,
                "created_at": now(),
                "updated_at": now(),
            },
        )
    update_row(
        "reviews_review_task",
        row_id,
        {"reviewer_id": user["id"], "reviewer_level": user.get("reviewer_level"), "updated_at": now()},
    )
    return {"message": "Review opinions submitted", "task_id": row_id, "review_result_id": result["id"]}


@app.get("/api/reviews/results/{row_id}/download_report/")
def download_report(row_id: int, user: dict[str, Any] = Depends(current_user)):
    return report_response(row_id, download=True)


@app.get("/api/reviews/results/{row_id}/preview_report/")
def preview_report(row_id: int, user: dict[str, Any] = Depends(current_user)):
    return report_response(row_id, download=False)


def report_response(row_id: int, download: bool):
    result = db_one("SELECT * FROM reviews_review_result WHERE id = %s", (row_id,))
    if not result or not result.get("report_path"):
        raise HTTPException(status_code=404, detail="Report not found")
    path = MEDIA_ROOT / str(result["report_path"])
    if not path.exists():
        raise HTTPException(status_code=404, detail="Report not found")
    return FileResponse(path, filename=path.name if download else None)


@app.post("/api/reviews/cycles/summarize_opinions/")
async def summarize_opinions(request: Request, user: dict[str, Any] = Depends(current_user)):
    payload = await request.json()
    contract_id = payload.get("contract_id")
    if not contract_id:
        raise HTTPException(status_code=400, detail="contract_id is required")
    opinions = db_query(
        """
        SELECT o.*, u.reviewer_level, u.username AS reviewer_name
        FROM reviews_review_opinion o
        JOIN reviews_review_result r ON r.id = o.review_result_id
        LEFT JOIN users_user u ON u.id = o.reviewer_id
        WHERE r.contract_id = %s AND COALESCE(o.is_deleted, 0) = 0
        ORDER BY o.created_at DESC
        """,
        (contract_id,),
    )
    summary_table = build_summary_table(opinions)
    return {"success": True, "summary_table": summary_table}


def build_summary_table(opinions: list[dict[str, Any]]) -> dict[str, Any]:
    def item(opinion: dict[str, Any]) -> dict[str, Any]:
        return {
            "reviewer_level": opinion.get("reviewer_level") or "",
            "reviewer_name": opinion.get("reviewer_name") or "",
            "type": opinion.get("opinion_type") or "",
            "risk_level": opinion.get("risk_level") or "",
            "content": opinion.get("opinion_content") or "",
            "suggestion": opinion.get("suggestion") or "",
        }

    items = [item(opinion) for opinion in opinions]
    return {
        "statistics": {
            "total_opinions": len(items),
            "high_risk_count": sum(1 for opinion in opinions if opinion.get("risk_level") == "high"),
            "medium_risk_count": sum(1 for opinion in opinions if opinion.get("risk_level") == "medium"),
            "low_risk_count": sum(1 for opinion in opinions if opinion.get("risk_level") == "low"),
        },
        "level1_opinions": [it for it in items if it["reviewer_level"] == "level1"],
        "level2_opinions": [it for it in items if it["reviewer_level"] == "level2"],
        "level3_opinions": [it for it in items if it["reviewer_level"] == "level3"],
        "all_opinions": items,
    }


@app.post("/api/reviews/cycles/feedback_to_drafter/")
async def feedback_to_drafter(request: Request, user: dict[str, Any] = Depends(current_user)):
    payload = await request.json()
    return {"success": True, "message": "Feedback recorded", "summary_table": payload.get("summary_table")}


@app.post("/api/reviews/cycles/resubmit_for_review/")
async def resubmit_for_review(request: Request, user: dict[str, Any] = Depends(current_user)):
    payload = await request.json()
    contract_id = payload.get("contract_id")
    if not contract_id:
        raise HTTPException(status_code=400, detail="contract_id is required")
    task = insert_row(
        "reviews_review_task",
        {
            "contract_id": contract_id,
            "task_type": "auto",
            "status": "pending",
            "priority": 0,
            "celery_task_id": "",
            "error_message": "",
            "created_by_id": user["id"],
            "created_at": now(),
            "updated_at": now(),
        },
    )
    return {"success": True, "message": "Resubmitted", "review_task_id": task["id"]}


@app.post("/api/reviews/ai-model-configs/{row_id}/set_default/")
def set_default_ai_config(row_id: int, user: dict[str, Any] = Depends(current_user)):
    db_execute("UPDATE reviews_ai_model_config SET is_default = 0")
    update_row("reviews_ai_model_config", row_id, {"is_default": 1, "is_active": 1, "updated_at": now()})
    return {"message": "Default config updated", "config": get_resource(SPECS["ai_model_configs"], row_id)}


@app.post("/api/reviews/ai-model-configs/{row_id}/test_connection/")
def test_ai_connection(row_id: int, user: dict[str, Any] = Depends(current_user)):
    config = get_resource(SPECS["ai_model_configs"], row_id)
    if not config.get("api_key"):
        return {"success": False, "message": "API key is empty"}
    return {"success": True, "message": "Config is valid", "model": config.get("default_model")}


@app.exception_handler(MySQLdb.Error)
async def mysql_error_handler(request: Request, exc: MySQLdb.Error):
    return PlainTextResponse(f"Database error: {exc}", status_code=500)
